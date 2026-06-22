"""Replace appendix 'کد اجرایی الگوریتم (پایتون)' with updated code export."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


DOCX_PATH = Path(r"e:\projects\thesis_project_v2\main_updated.docx")
CODE_PATH = Path(r"e:\projects\thesis_project_v2\code_export.py.py")


def _insert_paragraph_after(paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    if text:
        new_para.add_run(text)
    return new_para


def _delete_paragraph(paragraph) -> None:
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None  # type: ignore[attr-defined]


def _read_code_for_appendix(py_path: Path) -> str:
    text = py_path.read_text(encoding="utf-8", errors="replace")
    # Keep it as a clean .py excerpt (drop nbconvert header noise a bit).
    lines = text.splitlines()
    # Remove leading shebang/encoding and empty leading comment blocks.
    while lines and (lines[0].startswith("#!") or "coding:" in lines[0] or lines[0].strip() == ""):
        lines.pop(0)
    # Drop leading markdown-comment banner if present (starts with "# # ").
    while lines and lines[0].lstrip().startswith("#"):
        # stop once we hit first real import/def
        nxt = lines[0].lstrip("#").strip()
        if nxt.startswith("In[") or nxt == "" or nxt.startswith("رمزنگاری") or nxt.startswith("##") or nxt.startswith("#"):
            lines.pop(0)
            continue
        break
    # Hard trim: keep from first import.
    for i, ln in enumerate(lines):
        if ln.startswith("import ") or ln.startswith("from "):
            lines = lines[i:]
            break
    # Ensure trailing newline.
    return "\n".join(lines).rstrip() + "\n"


def main() -> None:
    if not DOCX_PATH.exists():
        raise FileNotFoundError(DOCX_PATH)
    if not CODE_PATH.exists():
        raise FileNotFoundError(CODE_PATH)

    code_text = _read_code_for_appendix(CODE_PATH)
    doc = Document(DOCX_PATH)

    # Find appendix heading line: "کد اجرایی الگوریتم (پایتون)"
    heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "کد اجرایی الگوریتم (پایتون)":
            heading_idx = i
            break
    if heading_idx is None:
        raise RuntimeError("Could not find appendix heading 'کد اجرایی الگوریتم (پایتون)'")

    # Determine the code block region: from first code line after heading until next non-code section.
    # In this document, code starts immediately after heading, and continues until paragraph index ~1340.
    start = heading_idx + 1
    # Heuristic end: first paragraph after start that is clearly NOT code and not empty, after a large run.
    end = None
    for j in range(start, len(doc.paragraphs)):
        t = doc.paragraphs[j].text
        if j > start + 20 and t.strip().startswith("پیوست"):
            end = j
            break
    if end is None:
        # fallback: stop when we hit a heading-like paragraph (very rare in this appendix)
        for j in range(start, len(doc.paragraphs)):
            style = doc.paragraphs[j].style.name if doc.paragraphs[j].style else ""
            if j > start + 20 and ("Heading" in style or style.startswith("toc")):
                end = j
                break
    if end is None:
        end = len(doc.paragraphs)

    # Remove existing code paragraphs.
    # Delete from end-1 down to start to keep indices valid.
    deleted = 0
    for j in range(end - 1, start - 1, -1):
        _delete_paragraph(doc.paragraphs[j])
        deleted += 1

    # Insert new code as multiple paragraphs (one per line) to preserve monospace wrapping.
    # Insert after heading paragraph.
    heading_para = doc.paragraphs[heading_idx]
    cur = heading_para

    for ln in code_text.splitlines():
        p = _insert_paragraph_after(cur, ln, style="Normal")
        # Force monospace where possible.
        if p.runs:
            r = p.runs[0]
            try:
                r.font.name = "Courier New"
                r.font.size = None  # keep document default if any
            except Exception:
                pass
        cur = p

    doc.save(DOCX_PATH)
    print(f"Updated appendix code: deleted {deleted} old paragraphs, inserted {len(code_text.splitlines())} lines.")


if __name__ == "__main__":
    main()

