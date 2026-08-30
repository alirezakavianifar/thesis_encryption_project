import os
import shutil
import re
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from projectcfg import CFG

def create_element(name):
    return OxmlElement(name)

def add_page_break(p):
    # Remove existing runs from paragraph
    for child in list(p._p):
        if child.tag.endswith(('r', 'hyperlink', 'fldSimple')):
            p._p.remove(child)
    r = create_element('w:r')
    br = create_element('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p._p.append(r)

def set_p_rtl(p, align="both"):
    pPr = p._p.get_or_add_pPr()
    
    # Enable paragraph-level bidi (RTL)
    bidi = pPr.find(qn('w:bidi'))
    if bidi is None:
        bidi = create_element('w:bidi')
        pPr.append(bidi)
    bidi.set(qn('w:val'), '1')

    # Set paragraph alignment (both/justify, center, right, etc.)
    if align:
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = create_element('w:jc')
            pPr.append(jc)
        jc.set(qn('w:val'), align)

def is_in_math(element, stop_element=None):
    cur = element.getparent()
    while cur is not None and cur != stop_element:
        tag = cur.tag
        if isinstance(tag, str) and ('math' in tag.lower() or tag.endswith(('oMath', 'oMathPara', 'r', 'ctrlPr'))):
            if 'http://schemas.openxmlformats.org/officeDocument/2006/math' in tag:
                return True
        cur = cur.getparent()
    return False

def apply_run_rtl_font(r, cs_font=CFG.font_family_persian, ascii_font=CFG.font_family_latin):
    r_el = r._r if hasattr(r, '_r') else r
    rPr = r_el.get_or_add_rPr()
    
    # 0. If run contains a footnoteReference or footnoteRef, enforce superscript & 9pt size
    if r_el.find(qn('w:footnoteReference')) is not None or r_el.find(qn('w:footnoteRef')) is not None:
        vert = rPr.find(qn('w:vertAlign'))
        if vert is None:
            vert = create_element('w:vertAlign')
            rPr.append(vert)
        vert.set(qn('w:val'), 'superscript')
        sz = rPr.find(qn('w:sz'))
        if sz is None:
            sz = create_element('w:sz')
            rPr.append(sz)
        sz.set(qn('w:val'), '18')
        szCs = rPr.find(qn('w:szCs'))
        if szCs is None:
            szCs = create_element('w:szCs')
            rPr.append(szCs)
        szCs.set(qn('w:val'), '18')

    # 1. Enable run-level RTL (Crucial for Persian font selection & correct colon placement)
    rtl = rPr.find(qn('w:rtl'))
    if rtl is None:
        rtl = create_element('w:rtl')
        rPr.append(rtl)
    rtl.set(qn('w:val'), '1')

    lang = rPr.find(qn('w:lang'))
    if lang is None:
        lang = create_element('w:lang')
        rPr.append(lang)
    lang.set(qn('w:bidi'), 'fa-IR')

    # 2. Assign Complex Script (cs) font for Persian and ASCII font for Latin
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = create_element('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:cs'), cs_font)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)

    # 3. Synchronize bold tags (w:b and w:bCs for Persian)
    b = rPr.find(qn('w:b'))
    if b is not None:
        bCs = rPr.find(qn('w:bCs'))
        if bCs is None:
            rPr.append(create_element('w:bCs'))

    # 4. Clean residual blue / theme colors from runs to ensure standard black text
    color = rPr.find(qn('w:color'))
    if color is not None:
        if qn('w:themeColor') in color.attrib:
            del color.attrib[qn('w:themeColor')]
        if qn('w:themeShade') in color.attrib:
            del color.attrib[qn('w:themeShade')]
        if color.get(qn('w:val')) in ['003366', '004682', '365F91', '4F81BD', '243F60', '1F497D']:
            color.set(qn('w:val'), '000000')

def _inject_page_numbering(docx_path, persian_font, latin_font):
    """
    Injects two-section page numbering into a DOCX via direct ZIP manipulation:
      Section 1 (front matter, up to Chapter 1):
        - Arabic Abjad letters (الف، ب، ج...) in footer
        - First page (title page) has blank footer (no number visible)
      Section 2 (body, from فصل اول onwards):
        - Persian/Eastern-Arabic numerals (۱، ۲، ۳...) in footer, restarting at 1
    """
    import zipfile, tempfile
    from lxml import etree

    W  = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    R  = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    PKG = 'http://schemas.openxmlformats.org/package/2006/relationships'
    CT  = 'http://schemas.openxmlformats.org/package/2006/content-types'
    FOOTER_TYPE = f'{R}/footer'
    FOOTER_MIME = 'application/vnd.openxmlformats-officedocument.wordprocessingml.footer+xml'

    def w(tag): return '{%s}%s' % (W, tag)
    def r(tag): return '{%s}%s' % (R, tag)

    # --- Footer XML builders ---
    def _footer_page_number(centered=True):
        """Footer with a PAGE field."""
        align = 'center' if centered else 'right'
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:p>'
            f'<w:pPr><w:bidi w:val="1"/><w:jc w:val="{align}"/></w:pPr>'
            f'<w:r><w:rPr><w:rFonts w:cs="{persian_font}" w:ascii="{latin_font}" w:hAnsi="{latin_font}"/></w:rPr>'
            '<w:fldChar w:fldCharType="begin"/></w:r>'
            f'<w:r><w:rPr><w:rFonts w:cs="{persian_font}" w:ascii="{latin_font}" w:hAnsi="{latin_font}"/></w:rPr>'
            '<w:instrText xml:space="preserve"> PAGE </w:instrText></w:r>'
            f'<w:r><w:rPr><w:rFonts w:cs="{persian_font}" w:ascii="{latin_font}" w:hAnsi="{latin_font}"/></w:rPr>'
            '<w:fldChar w:fldCharType="end"/></w:r>'
            '</w:p>'
            '</w:ftr>'
        ).encode('utf-8')

    def _footer_empty():
        """Empty footer (for title page first-page override)."""
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:ftr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:p><w:pPr><w:bidi w:val="1"/><w:jc w:val="center"/></w:pPr></w:p>'
            '</w:ftr>'
        ).encode('utf-8')

    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
    os.close(tmp_fd)

    with zipfile.ZipFile(docx_path, 'r') as zin:
        doc_bytes   = zin.read('word/document.xml')
        rels_bytes  = zin.read('word/_rels/document.xml.rels')
        ct_bytes    = zin.read('[Content_Types].xml')

        doc_root  = etree.fromstring(doc_bytes)
        rels_root = etree.fromstring(rels_bytes)
        ct_root   = etree.fromstring(ct_bytes)
        body      = doc_root.find(w('body'))

        # --- Allocate relationship IDs ---
        existing = [int(el.get('Id', 'rId0')[3:]) for el in rels_root
                    if el.get('Id', '').startswith('rId') and el.get('Id', 'rId0')[3:].isdigit()]
        nxt = max(existing, default=10) + 1
        rid_fm_default = f'rId{nxt}'      # front-matter regular footer (abjad)
        rid_fm_first   = f'rId{nxt+1}'    # front-matter first-page footer (blank)
        rid_body       = f'rId{nxt+2}'    # body footer (Persian nums)

        # --- Add relationships ---
        PKG_ns = 'http://schemas.openxmlformats.org/package/2006/relationships'
        for rid, fname in [
            (rid_fm_default, 'footer_fm.xml'),
            (rid_fm_first,   'footer_blank.xml'),
            (rid_body,       'footer_body.xml'),
        ]:
            el = etree.SubElement(rels_root, '{%s}Relationship' % PKG_ns)
            el.set('Id', rid); el.set('Type', FOOTER_TYPE); el.set('Target', fname)

        # --- Add content types ---
        for fname in ['footer_fm.xml', 'footer_blank.xml', 'footer_body.xml']:
            ov = etree.SubElement(ct_root, '{%s}Override' % CT)
            ov.set('PartName', f'/word/{fname}'); ov.set('ContentType', FOOTER_MIME)

        # --- Locate Chapter 1 heading and the paragraph just before it ---
        body_children = list(body)
        ch1_idx = None
        for idx, child in enumerate(body_children):
            if child.tag != w('p'): continue
            texts = ''.join(t.text or '' for t in child.iter(w('t')))
            if 'فصل اول' in texts or ('کلیات' in texts and 'تحقیق' in texts):
                # Confirm it is a Heading 1
                pPr = child.find(w('pPr'))
                if pPr is not None:
                    ps = pPr.find(w('pStyle'))
                    if ps is not None and 'Heading' in ps.get(w('val'), ''):
                        ch1_idx = idx
                        break

        if ch1_idx is None:
            print("Warning: Chapter 1 heading not found; page numbering not split into sections")
            shutil.move(tmp_path, docx_path)
            return

        # Remove any existing pPr/sectPr in intermediate paragraphs
        for child in body_children:
            if child.tag == w('p'):
                pPr = child.find(w('pPr'))
                if pPr is not None:
                    for sp in pPr.findall(w('sectPr')):
                        pPr.remove(sp)

        # Find last w:p before Chapter 1 heading to attach Section 1 sectPr
        prev_p = None
        for i in range(ch1_idx - 1, -1, -1):
            if body_children[i].tag == w('p'):
                prev_p = body_children[i]; break

        def make_sectPr(fmt, start, include_titlePg, default_rid, first_rid=None):
            sp = etree.Element(w('sectPr'))
            if include_titlePg and first_rid:
                fr = etree.SubElement(sp, w('footerReference'))
                fr.set(w('type'), 'first'); fr.set(r('id'), first_rid)
            dr = etree.SubElement(sp, w('footerReference'))
            dr.set(w('type'), 'default'); dr.set(r('id'), default_rid)
            pgSz = etree.SubElement(sp, w('pgSz'))
            pgSz.set(w('w'), '11906'); pgSz.set(w('h'), '16838')
            pgMar = etree.SubElement(sp, w('pgMar'))
            pgMar.set(w('top'), '1417'); pgMar.set(w('right'), '1701')
            pgMar.set(w('bottom'), '1417'); pgMar.set(w('left'), '1417')
            pgMar.set(w('footer'), '851')
            pn = etree.SubElement(sp, w('pgNumType'))
            pn.set(w('fmt'), fmt); pn.set(w('start'), str(start))
            if include_titlePg:
                etree.SubElement(sp, w('titlePg'))
            etree.SubElement(sp, w('bidi')).set(w('val'), '1')
            return sp

        # Attach front-matter sectPr to the paragraph before Chapter 1
        if prev_p is not None:
            pPr = prev_p.find(w('pPr'))
            if pPr is None:
                pPr = etree.Element(w('pPr')); prev_p.insert(0, pPr)
            pPr.append(make_sectPr(
                fmt='arabicAbjad', start=1,
                include_titlePg=True,
                default_rid=rid_fm_default,
                first_rid=rid_fm_first,
            ))

        # Update the final body sectPr
        main_sp = body.find(w('sectPr'))
        if main_sp is None:
            main_sp = etree.SubElement(body, w('sectPr'))
        for tag in ['footerReference', 'headerReference', 'pgNumType', 'titlePg', 'pgSz', 'pgMar']:
            for el in main_sp.findall(w(tag)): main_sp.remove(el)
        for child in list(make_sectPr(
            fmt='decimalArabic', start=1,
            include_titlePg=False,
            default_rid=rid_body,
        )):
            main_sp.append(child)

        # Serialise
        doc_out  = etree.tostring(doc_root,  encoding='utf-8', xml_declaration=True)
        rels_out = etree.tostring(rels_root, encoding='utf-8', xml_declaration=True)
        ct_out   = etree.tostring(ct_root,   encoding='utf-8', xml_declaration=True)

        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                name = item.filename
                if   name == 'word/document.xml':          zout.writestr(item, doc_out)
                elif name == 'word/_rels/document.xml.rels': zout.writestr(item, rels_out)
                elif name == '[Content_Types].xml':         zout.writestr(item, ct_out)
                else:                                       zout.writestr(item, zin.read(name))
            # Write new footer parts
            zout.writestr('word/footer_fm.xml',    _footer_page_number())   # abjad (auto by pgNumType)
            zout.writestr('word/footer_blank.xml', _footer_empty())          # title page = no number
            zout.writestr('word/footer_body.xml',  _footer_page_number())   # Persian nums (auto)

    shutil.move(tmp_path, docx_path)
    print("Injected dual-section page numbering (arabicAbjad front / decimalArabic body)")


def postprocess_doc():

    if not os.path.exists(CFG.intermediate_docx):
        raise FileNotFoundError(f"Intermediate docx missing: {CFG.intermediate_docx}")

    doc = Document(CFG.intermediate_docx)

    # Enable section-level bidi for all sections
    for section in doc.sections:
        sectPr = section._sectPr
        bidi = sectPr.find(qn('w:bidi'))
        if bidi is None:
            bidi = create_element('w:bidi')
            sectPr.append(bidi)
        bidi.set(qn('w:val'), '1')

    # Harvest headings, figures, tables for static TOC/LOT/LOF generation
    headings = []
    figures = []
    tables = []
    for p in doc.paragraphs:
        t_str = p.text.strip()
        if not t_str or t_str.startswith("@@"): continue
        sname = p.style.name.lower()
        if 'heading' in sname:
            headings.append((p.style.name, t_str))
        elif t_str.startswith("شکل"):
            figures.append(t_str)
        elif t_str.startswith("جدول"):
            tables.append(t_str)

    def build_toc_field(anchor_p, field_instr, entry_paragraphs):
        """Build a proper w:fldChar begin/instrText/sep/end TOC field.
        
        The entry_paragraphs list lives between the 'separate' and 'end' markers,
        which is exactly the display region Word repopulates when you press F9
        or click 'Update Table'. This makes the TOC fully updateable in Word.
        """
        # --- paragraph that holds: fldChar(begin) + instrText ---
        p_begin = create_element('w:p')
        pPr_b = create_element('w:pPr')
        bidi_b = create_element('w:bidi'); bidi_b.set(qn('w:val'), '1'); pPr_b.append(bidi_b)
        p_begin.append(pPr_b)

        r_begin = create_element('w:r')
        fc_begin = create_element('w:fldChar')
        fc_begin.set(qn('w:fldCharType'), 'begin')
        fc_begin.set(qn('w:dirty'), 'true')   # mark dirty so Word recalculates on open
        r_begin.append(fc_begin)
        p_begin.append(r_begin)

        r_instr = create_element('w:r')
        instr_el = create_element('w:instrText')
        instr_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        instr_el.text = ' ' + field_instr + ' '
        r_instr.append(instr_el)
        p_begin.append(r_instr)

        anchor_p._p.addnext(p_begin)
        prev = p_begin

        # --- paragraph that holds: fldChar(separate) ---
        p_sep = create_element('w:p')
        pPr_s = create_element('w:pPr')
        bidi_s = create_element('w:bidi'); bidi_s.set(qn('w:val'), '1'); pPr_s.append(bidi_s)
        p_sep.append(pPr_s)
        r_sep = create_element('w:r')
        fc_sep = create_element('w:fldChar'); fc_sep.set(qn('w:fldCharType'), 'separate')
        r_sep.append(fc_sep)
        p_sep.append(r_sep)
        prev.addnext(p_sep)
        prev = p_sep

        # --- static display paragraphs (between sep and end) ---
        for ep in entry_paragraphs:
            prev.addnext(ep)
            prev = ep

        # --- paragraph that holds: fldChar(end) ---
        p_end = create_element('w:p')
        pPr_e = create_element('w:pPr')
        bidi_e = create_element('w:bidi'); bidi_e.set(qn('w:val'), '1'); pPr_e.append(bidi_e)
        p_end.append(pPr_e)
        r_end = create_element('w:r')
        fc_end = create_element('w:fldChar'); fc_end.set(qn('w:fldCharType'), 'end')
        r_end.append(fc_end)
        p_end.append(r_end)
        prev.addnext(p_end)

        return p_end  # caller can chain after the end marker

    def create_toc_entry_with_dots(title, page_num_str, indent_dxa=0, is_bold=False):
        p_entry = create_element('w:p')
        pPr = create_element('w:pPr')
        bidi = create_element('w:bidi'); bidi.set(qn('w:val'), '1'); pPr.append(bidi)
        jc = create_element('w:jc'); jc.set(qn('w:val'), 'right'); pPr.append(jc)
        
        tabs = create_element('w:tabs')
        tab = create_element('w:tab')
        tab.set(qn('w:val'), 'left')
        tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'), '8500')
        tabs.append(tab)
        pPr.append(tabs)
        
        if indent_dxa > 0:
            ind = create_element('w:ind'); ind.set(qn('w:right'), str(indent_dxa)); pPr.append(ind)
        p_entry.append(pPr)
        
        r1 = create_element('w:r')
        rPr1 = create_element('w:rPr')
        rtl1 = create_element('w:rtl'); rtl1.set(qn('w:val'), '1'); rPr1.append(rtl1)
        rFonts1 = create_element('w:rFonts')
        rFonts1.set(qn('w:cs'), CFG.font_family_persian)
        rFonts1.set(qn('w:ascii'), CFG.font_family_latin)
        rFonts1.set(qn('w:hAnsi'), CFG.font_family_latin)
        rPr1.append(rFonts1)
        if is_bold:
            b1 = create_element('w:b'); rPr1.append(b1)
        r1.append(rPr1)
        t1 = create_element('w:t'); t1.text = title; r1.append(t1)
        p_entry.append(r1)
        
        r2 = create_element('w:r')
        rPr2 = create_element('w:rPr')
        rtl2 = create_element('w:rtl'); rtl2.set(qn('w:val'), '1'); rPr2.append(rtl2)
        r2.append(rPr2)
        r2.append(create_element('w:tab'))
        p_entry.append(r2)
        
        r3 = create_element('w:r')
        rPr3 = create_element('w:rPr')
        rtl3 = create_element('w:rtl'); rtl3.set(qn('w:val'), '1'); rPr3.append(rtl3)
        rFonts3 = create_element('w:rFonts')
        rFonts3.set(qn('w:cs'), CFG.font_family_persian)
        rFonts3.set(qn('w:ascii'), CFG.font_family_latin)
        rFonts3.set(qn('w:hAnsi'), CFG.font_family_latin)
        rPr3.append(rFonts3)
        if is_bold:
            b3 = create_element('w:b'); rPr3.append(b3)
        r3.append(rPr3)
        t3 = create_element('w:t'); t3.text = page_num_str; r3.append(t3)
        p_entry.append(r3)
        
        return p_entry

    def to_fa(n):
        return str(n).translate(str.maketrans('0123456789', '۰۱۲۳۴۵۶۷۸۹'))

    # Arabic Abjad letter sequence for front-matter page numbering
    # Order: الف=1, ب=2, ج=3, د=4, هـ=5, و=6, ز=7, ح=8, ط=9, ي=10, ك=11, ل=12, م=13, ن=14, س=15, ع=16, ف=17, ص=18, ق=19, ر=20, ش=21, ت=22, ث=23, خ=24, ذ=25, ض=26, ظ=27, غ=28
    ABJAD = ['الف','ب','ج','د','هـ','و','ز','ح','ط','ي','ك','ل','م','ن','س','ع','ف','ص','ق','ر','ش','ت','ث','خ','ذ','ض','ظ','غ']

    def to_abjad(n):
        """Convert 1-based integer to Arabic Abjad letter string."""
        if 1 <= n <= len(ABJAD):
            return ABJAD[n - 1]
        return str(n)  # fallback for very long front-matters

    # Front-matter sequential page layout:
    # The title section has 2 pages (Persian title + English title), both with blank footer.
    # Every preliminary item after that occupies exactly 1 physical page, EXCEPT the TOC
    # which spans as many pages as needed. We assign Abjad numbers sequentially.
    #
    # Physical page:  1=الف  2=ب  [both title pages - blank footer, not shown]
    # Page 3=ج  : صورتجلسه دفاع
    # Page 4=د  : تعهدنامه اصالت
    # Page 5=هـ  : منشور اخلاق پژوهش
    # Page 6=و  : تقدیم
    # Page 7=ز  : سپاسگزاری
    # Page 8=ح  : فهرست مطالب (TOC - may span several pages, heading on page 8)
    # Page after TOC: فهرست جداول (م=13 if TOC is 5 pages)
    # ... then فهرست اشکال, فهرست اختصارات, چکیده فارسی
    #
    # The exact page numbers for TOC sub-items and post-TOC front-matter are left as '?' 
    # because Word will fill them in correctly on F9 update (the field knows the actual Abjad
    # page numbers from the section's pgNumType).
    FRONT_MATTER_ABJAD = {
        'صورتجلسه دفاع از پایان‌نامه': to_abjad(3),
        'تعهدنامه اصالت رساله یا پایان‌نامه تحصیلی': to_abjad(4),
        'منشور اخلاق پژوهش': to_abjad(5),
        'تقدیم': to_abjad(6),
        'سپاسگزاری': to_abjad(7),
        'فهرست مطالب': to_abjad(8),
        'فهرست جداول': to_abjad(13),   # assuming TOC spans pages 8-12
        'فهرست اشکال': to_abjad(14),
        'فهرست اختصارات و علائم': to_abjad(15),
        'چکیده': to_abjad(16),
    }

    # Body front-matter keywords to detect (items that appear in TOC but are in the front section)
    FRONT_MATTER_KEYWORDS = ['صورتجلسه','تعهدنامه','منشور','تقدیم','سپاسگزاری',
                            'فهرست مطالب','فهرست جداول','فهرست اشکال',
                            'فهرست اختصارات','چکیده']

    def is_front_matter(title_text):
        return any(kw in title_text for kw in FRONT_MATTER_KEYWORDS)

    # Dynamic harvesting of exact page numbers from exported PDF preview if available
    dynamic_page_map = {}
    pdf_preview_path = os.path.abspath(os.path.join(os.path.dirname(CFG.output_docx), "thesis_preview.pdf"))
    if os.path.exists(pdf_preview_path):
        import fitz
        try:
            pdf_doc = fitz.open(pdf_preview_path)
            norm_pages = [re.sub(r'[^\u0600-\u06FF\u06f0-\u06f9a-zA-Z0-9]', '', pdf_doc[pno].get_text()) for pno in range(17, len(pdf_doc))]
            for h_style, h_text in headings:
                if is_front_matter(h_text): continue  # skip front-matter; handled by FRONT_MATTER_ABJAD
                target = re.sub(r'[^\u0600-\u06FF\u06f0-\u06f9a-zA-Z0-9]', '', h_text)
                if len(target) < 3: continue
                search_stem = target[:10] if len(target) >= 10 else target
                for p_idx, p_norm in enumerate(norm_pages):
                    if search_stem in p_norm:
                        dynamic_page_map[h_text] = to_fa(max(1, p_idx + 1))
                        break
        except Exception as e:
            print("PDF page harvest notice:", e)

    def get_pg_num(title_text):
        """Return the display page number string for a heading.
        Front-matter headings get Arabic Abjad letters (الف، ب، ج...).
        Body headings get Persian decimal numerals (۱، ۲، ۳...) from PDF harvest.
        """
        # 1. Front-matter exact match
        for k, v in FRONT_MATTER_ABJAD.items():
            if k in title_text or title_text in k:
                return v
        # 2. General front-matter keyword match -> Abjad placeholder
        if is_front_matter(title_text):
            return '...'   # Word will fill in on F9
        # 3. Body heading: use PDF-harvested Persian decimal
        if title_text in dynamic_page_map:
            return dynamic_page_map[title_text]
        # Fuzzy match against dynamic map
        target = re.sub(r'[^\u0600-\u06FF\u06f0-\u06f9a-zA-Z0-9]', '', title_text)
        search_stem = target[:10] if len(target) >= 10 else target
        for k, v in dynamic_page_map.items():
            k_norm = re.sub(r'[^\u0600-\u06FF\u06f0-\u06f9a-zA-Z0-9]', '', k)
            if search_stem in k_norm or k_norm[:10] in target:
                return v
        return '۱'

    is_centered = False

    for p in list(doc.paragraphs):
        txt = p.text.strip()

        # Handle sentinel page breaks (@@PB@@)
        if txt == "@@PB@@":
            add_page_break(p)
            continue

        # Handle TOC sentinel (@@TOC@@)
        if txt == "@@TOC@@":
            p.text = "فهرست مطالب"
            p.style = "Heading 1"
            set_p_rtl(p, align="right")
            for r in p.runs: apply_run_rtl_font(r)
            toc_entries = []
            for h_style, h_text in headings:
                if any(x in h_text for x in ["فهرست", "صورت‌جلسه", "تعهدنامه", "منشور"]):
                    continue
                ind_val = 0
                if "heading 2" in h_style.lower(): ind_val = 360
                elif "heading 3" in h_style.lower(): ind_val = 720
                pg_num = get_pg_num(h_text)
                toc_entries.append(create_toc_entry_with_dots(h_text, pg_num, indent_dxa=ind_val, is_bold=("heading 1" in h_style.lower())))
            build_toc_field(p, 'TOC \\o "1-3" \\h \\z \\u', toc_entries)
            continue

        # Handle LOT sentinel (@@LOT@@)
        if txt == "@@LOT@@":
            p.text = "فهرست جداول"
            p.style = "Heading 1"
            set_p_rtl(p, align="right")
            for r in p.runs: apply_run_rtl_font(r)
            lot_entries = []
            for t_caption in tables:
                pg_num = get_pg_num(t_caption)
                lot_entries.append(create_toc_entry_with_dots(t_caption, pg_num, indent_dxa=0, is_bold=False))
            build_toc_field(p, 'TOC \\t "CaptionTable,1" \\h \\z', lot_entries)
            continue

        # Handle LOF sentinel (@@LOF@@)
        if txt == "@@LOF@@":
            p.text = "فهرست اشکال"
            p.style = "Heading 1"
            set_p_rtl(p, align="right")
            for r in p.runs: apply_run_rtl_font(r)
            lof_entries = []
            for f_caption in figures:
                pg_num = get_pg_num(f_caption)
                lof_entries.append(create_toc_entry_with_dots(f_caption, pg_num, indent_dxa=0, is_bold=False))
            build_toc_field(p, 'TOC \\t "CaptionFigure,1" \\h \\z', lof_entries)
            continue

        # Handle sentinel centering (@@CENTER_START@@ / @@CENTER_END@@)
        if txt == "@@CENTER_START@@":
            is_centered = True
            p.text = ""
            continue
        elif txt == "@@CENTER_END@@":
            is_centered = False
            p.text = ""
            continue

        # Determine alignment and set outline level for native dynamic TOC building in Word
        st_name = p.style.name.lower()
        pPr = p._p.get_or_add_pPr()
        if "heading 1" in st_name:
            align_val = "right"
            ol = pPr.find(qn('w:outlineLvl'))
            if ol is None: ol = create_element('w:outlineLvl'); pPr.append(ol)
            ol.set(qn('w:val'), '0')
        elif "heading 2" in st_name:
            align_val = "right"
            ol = pPr.find(qn('w:outlineLvl'))
            if ol is None: ol = create_element('w:outlineLvl'); pPr.append(ol)
            ol.set(qn('w:val'), '1')
        elif "heading 3" in st_name:
            align_val = "right"
            ol = pPr.find(qn('w:outlineLvl'))
            if ol is None: ol = create_element('w:outlineLvl'); pPr.append(ol)
            ol.set(qn('w:val'), '2')
        elif "heading 4" in st_name:
            align_val = "right"
            ol = pPr.find(qn('w:outlineLvl'))
            if ol is None: ol = create_element('w:outlineLvl'); pPr.append(ol)
            ol.set(qn('w:val'), '3')
        elif is_centered or txt.startswith("شکل") or txt.startswith("جدول"):
            align_val = "center"
        else:
            align_val = "both"

        # Fix BiDi ordering on Heading paragraphs (ensuring section numbers and embedded Latin words are placed correctly RTL)
        if any(h in st_name for h in ["heading 1", "heading 2", "heading 3", "heading 4"]):
            for r_el in p._p.iter(qn('w:r')):
                for t in r_el.iter(qn('w:t')):
                    if t.text:
                        m_sec = re.match(r"^([\u200e\u200f]?[\u06f0-\u06f90-9]+(?:[\-\u200e\u200f][\u06f0-\u06f90-9]+)+[\u200e\u200f]?)\s+(.*)", t.text)
                        if m_sec:
                            raw_num = re.sub(r"[\u200e\u200f]", "", m_sec.group(1))
                            parts = raw_num.split('-')
                            bidi_num = "-".join(f"\u200f{part}\u200f" for part in parts)
                            rest = m_sec.group(2)
                            rest_clean = re.sub(r"(?<=\s)([A-Za-z0-9\-]+)(?=\s|$)", chr(0x200e) + r"\1" + chr(0x200f), rest)
                            t.text = f"{bidi_num} {rest_clean}"

        # Apply CaptionTable / CaptionFigure styles so that LOT/LOF TOC fields
        # can collect them via TOC \t "CaptionTable,1" / "CaptionFigure,1"
        clean_txt = re.sub(r"[\u200e\u200f\s]+", " ", txt).strip()
        if clean_txt.startswith("جدول"):
            try:
                p.style = "CaptionTable"
            except Exception:
                pass
        elif clean_txt.startswith("شکل"):
            try:
                p.style = "CaptionFigure"
            except Exception:
                pass

        # Handle SourceCode / Verbatim paragraphs (Python code listing in appendix)
        st_name = p.style.name.lower()
        if 'sourcecode' in st_name or 'verbatim' in st_name or p.style.name == 'SourceCode':
            pPr = p._p.get_or_add_pPr()
            bidi = pPr.find(qn('w:bidi'))
            if bidi is None:
                bidi = create_element('w:bidi')
                pPr.append(bidi)
            bidi.set(qn('w:val'), '0')
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = create_element('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'left')
            for r_el in p._p.iter(qn('w:r')):
                rPr = r_el.get_or_add_rPr()
                rtl = rPr.find(qn('w:rtl'))
                if rtl is not None:
                    rtl.set(qn('w:val'), '0')
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = create_element('w:rFonts')
                    rPr.append(rFonts)
                rFonts.set(qn('w:ascii'), 'Consolas')
                rFonts.set(qn('w:hAnsi'), 'Consolas')
                rFonts.set(qn('w:cs'), 'Consolas')
                sz = rPr.find(qn('w:sz'))
                if sz is None:
                    sz = create_element('w:sz')
                    rPr.append(sz)
                sz.set(qn('w:val'), '17')
                szCs = rPr.find(qn('w:szCs'))
                if szCs is None:
                    szCs = create_element('w:szCs')
                    rPr.append(szCs)
                szCs.set(qn('w:val'), '17')
            continue

        set_p_rtl(p, align=align_val)

        # Convert paragraph indents from left to right for RTL list wrapping
        pPr = p._p.get_or_add_pPr()
        if pPr.find(qn('w:numPr')) is not None or "list" in p.style.name.lower():
            ind = pPr.find(qn('w:ind'))
            if ind is not None:
                left_val = ind.get(qn('w:left'))
                if left_val:
                    ind.set(qn('w:right'), left_val)
                    if qn('w:left') in ind.attrib:
                        del ind.attrib[qn('w:left')]

        # Check if paragraph is a pure English reference entry
        has_persian = any('\u0600' <= ch <= '\u06ff' or '\u06f0' <= ch <= '\u06f9' for ch in txt)
        first_char = txt[0] if txt else ''
        is_english_text = first_char.isalpha() and ord(first_char) < 128
        
        if not has_persian and is_english_text and not is_centered and ("(" in txt or len(txt) > 80):
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = create_element('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), 'left')
            bidi = pPr.find(qn('w:bidi'))
            if bidi is not None:
                bidi.set(qn('w:val'), '0')

        # Apply run-level RTL & B Lotus font (iterating all XML runs, skipping math)
        for r_el in p._p.iter(qn('w:r')):
            if not is_in_math(r_el, p._p):
                apply_run_rtl_font(r_el)

    # Format tables: enable grid borders, center alignment, cell padding, and RTL fonts
    for table in doc.tables:
        try:
            table.style = 'Table Grid'
        except Exception:
            pass
        
        # Center table on page
        tblPr = table._tbl.tblPr
        jc = tblPr.find(qn('w:jc'))
        if jc is None:
            jc = create_element('w:jc')
            tblPr.append(jc)
        jc.set(qn('w:val'), 'center')
        
        # Format table borders
        borders = tblPr.find(qn('w:tblBorders'))
        if borders is None:
            borders = create_element('w:tblBorders')
            tblPr.append(borders)
        for border_name, sz in [('top', '8'), ('bottom', '8'), ('left', '6'), ('right', '6'), ('insideH', '4'), ('insideV', '4')]:
            b = borders.find(qn(f'w:{border_name}'))
            if b is None:
                b = create_element(f'w:{border_name}')
                borders.append(b)
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), sz)
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')

        # Cell margins (padding)
        cellMar = tblPr.find(qn('w:tblCellMar'))
        if cellMar is None:
            cellMar = create_element('w:tblCellMar')
            tblPr.append(cellMar)
        for m_name, val in [('top', '100'), ('bottom', '100'), ('left', '140'), ('right', '140')]:
            m = cellMar.find(qn(f'w:{m_name}'))
            if m is None:
                m = create_element(f'w:{m_name}')
                cellMar.append(m)
            m.set(qn('w:w'), val)
            m.set(qn('w:type'), 'dxa')

        # Format header row and body rows
        for r_idx, row in enumerate(table.rows):
            trPr = row._tr.get_or_add_trPr()
            cantSplit = trPr.find(qn('w:cantSplit'))
            if cantSplit is None:
                cantSplit = create_element('w:cantSplit')
                trPr.append(cantSplit)
            
            if r_idx == 0:
                tblHeader = trPr.find(qn('w:tblHeader'))
                if tblHeader is None:
                    tblHeader = create_element('w:tblHeader')
                    trPr.append(tblHeader)

            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                vAlign = tcPr.find(qn('w:vAlign'))
                if vAlign is None:
                    vAlign = create_element('w:vAlign')
                    tcPr.append(vAlign)
                vAlign.set(qn('w:val'), 'center')

                for p in cell.paragraphs:
                    set_p_rtl(p, align="center")
                    for r_el in p._p.iter(qn('w:r')):
                        if not is_in_math(r_el, p._p):
                            apply_run_rtl_font(r_el)
                            if r_idx == 0:
                                rPr = r_el.get_or_add_rPr()
                                b = rPr.find(qn('w:b'))
                                if b is None:
                                    b = create_element('w:b')
                                    rPr.append(b)
                                bCs = rPr.find(qn('w:bCs'))
                                if bCs is None:
                                    bCs = create_element('w:bCs')
                                    rPr.append(bCs)

    # Format native Word footnotes XML part
    for part in doc.part.package.parts:
        if 'footnote' in part.partname:
            from docx.oxml import parse_xml
            root = parse_xml(part.blob)
            for fn in root.iter(qn('w:footnote')):
                fn_id = fn.get(qn('w:id'))
                if fn_id in ['0', '-1']:
                    continue
                for p in fn.iter(qn('w:p')):
                    pPr = p.find(qn('w:pPr'))
                    if pPr is None:
                        pPr = create_element('w:pPr')
                        p.append(pPr)
                    
                    # Ensure FootnoteText style is referenced
                    pStyle = pPr.find(qn('w:pStyle'))
                    if pStyle is None:
                        pStyle = create_element('w:pStyle')
                        pPr.append(pStyle)
                    pStyle.set(qn('w:val'), 'FootnoteText')

                    # Gather full text of the footnote (excluding footnoteRef)
                    full_p_text = ''.join(t.text or '' for t in p.iter(qn('w:t'))).strip()
                    has_persian = any('\u0600' <= ch <= '\u06ff' or '\u06f0' <= ch <= '\u06f9' for ch in full_p_text)

                    if not has_persian:
                        # Pure Latin footnote (e.g. "Advanced Encryption Standard (AES)") -> LTR left-aligned
                        bidi = pPr.find(qn('w:bidi'))
                        if bidi is None:
                            bidi = create_element('w:bidi')
                            pPr.append(bidi)
                        bidi.set(qn('w:val'), '0')
                        jc = pPr.find(qn('w:jc'))
                        if jc is None:
                            jc = create_element('w:jc')
                            pPr.append(jc)
                        jc.set(qn('w:val'), 'left')

                        for r in p.iter(qn('w:r')):
                            rPr = r.find(qn('w:rPr'))
                            if rPr is None:
                                rPr = create_element('w:rPr')
                                r.append(rPr)
                            
                            is_fn_ref = r.find(qn('w:footnoteRef')) is not None
                            if is_fn_ref:
                                vert = rPr.find(qn('w:vertAlign'))
                                if vert is None:
                                    vert = create_element('w:vertAlign')
                                    rPr.append(vert)
                                vert.set(qn('w:val'), 'superscript')
                                sz = rPr.find(qn('w:sz'))
                                if sz is None:
                                    sz = create_element('w:sz')
                                    rPr.append(sz)
                                sz.set(qn('w:val'), '18')
                            else:
                                sz = rPr.find(qn('w:sz'))
                                if sz is None:
                                    sz = create_element('w:sz')
                                    rPr.append(sz)
                                sz.set(qn('w:val'), '20')

                            rtl = rPr.find(qn('w:rtl'))
                            if rtl is not None:
                                rtl.set(qn('w:val'), '0')
                            
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = create_element('w:rFonts')
                                rPr.append(rFonts)
                            rFonts.set(qn('w:ascii'), CFG.font_family_latin)
                            rFonts.set(qn('w:hAnsi'), CFG.font_family_latin)
                            rFonts.set(qn('w:cs'), CFG.font_family_persian)
                    else:
                        # Persian or Mixed footnote (e.g. "Differential Cryptanalysis: روشی...") -> RTL right-aligned
                        bidi = pPr.find(qn('w:bidi'))
                        if bidi is None:
                            bidi = create_element('w:bidi')
                            pPr.append(bidi)
                        bidi.set(qn('w:val'), '1')
                        jc = pPr.find(qn('w:jc'))
                        if jc is None:
                            jc = create_element('w:jc')
                            pPr.append(jc)
                        jc.set(qn('w:val'), 'right')

                        for r in p.iter(qn('w:r')):
                            rPr = r.find(qn('w:rPr'))
                            if rPr is None:
                                rPr = create_element('w:rPr')
                                r.append(rPr)

                            is_fn_ref = r.find(qn('w:footnoteRef')) is not None
                            if is_fn_ref:
                                vert = rPr.find(qn('w:vertAlign'))
                                if vert is None:
                                    vert = create_element('w:vertAlign')
                                    rPr.append(vert)
                                vert.set(qn('w:val'), 'superscript')
                                sz = rPr.find(qn('w:sz'))
                                if sz is None:
                                    sz = create_element('w:sz')
                                    rPr.append(sz)
                                sz.set(qn('w:val'), '18')
                                szCs = rPr.find(qn('w:szCs'))
                                if szCs is None:
                                    szCs = create_element('w:szCs')
                                    rPr.append(szCs)
                                szCs.set(qn('w:val'), '18')
                            else:
                                sz = rPr.find(qn('w:sz'))
                                if sz is None:
                                    sz = create_element('w:sz')
                                    rPr.append(sz)
                                sz.set(qn('w:val'), '20')
                                szCs = rPr.find(qn('w:szCs'))
                                if szCs is None:
                                    szCs = create_element('w:szCs')
                                    rPr.append(szCs)
                                szCs.set(qn('w:val'), '20')

                            rtl = rPr.find(qn('w:rtl'))
                            if rtl is None:
                                rtl = create_element('w:rtl')
                                rPr.append(rtl)
                            rtl.set(qn('w:val'), '1')

                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = create_element('w:rFonts')
                                rPr.append(rFonts)
                            rFonts.set(qn('w:cs'), CFG.font_family_persian)
                            rFonts.set(qn('w:ascii'), CFG.font_family_latin)
                            rFonts.set(qn('w:hAnsi'), CFG.font_family_latin)

                            # Handle mixed English term with Persian explanation like:
                            # "Differential Cryptanalysis: روشی..."
                            for t in r.iter(qn('w:t')):
                                if t.text:
                                    m_lead = re.match(r"^([A-Za-z0-9\s\-]+):\s*([\u0600-\u06FF].*)", t.text)
                                    if m_lead:
                                        eng_term = m_lead.group(1).strip()
                                        fa_rest = m_lead.group(2)
                                        t.text = f"\u200e{eng_term}\u200f: {fa_rest}"
            from lxml import etree
            part._blob = etree.tostring(root, encoding='utf-8', xml_declaration=True)

    # Format numbering.xml for native RTL numbered and bullet lists
    for part in doc.part.package.parts:
        if 'numbering' in part.partname:
            from docx.oxml import parse_xml
            root = parse_xml(part.blob)
            for lvl in root.iter(qn('w:lvl')):
                ilvl = int(lvl.get(qn('w:ilvl'), '0'))
                lvlJc = lvl.find(qn('w:lvlJc'))
                if lvlJc is None:
                    lvlJc = create_element('w:lvlJc')
                    lvl.append(lvlJc)
                lvlJc.set(qn('w:val'), 'right')
                
                pPr = lvl.find(qn('w:pPr'))
                if pPr is None:
                    pPr = create_element('w:pPr')
                    lvl.append(pPr)
                bidi = pPr.find(qn('w:bidi'))
                if bidi is None:
                    bidi = create_element('w:bidi')
                    pPr.append(bidi)
                bidi.set(qn('w:val'), '1')

                ind = pPr.find(qn('w:ind'))
                if ind is None:
                    ind = create_element('w:ind')
                    pPr.append(ind)
                right_val = 720 * (ilvl + 1)
                ind.set(qn('w:right'), str(right_val))
                ind.set(qn('w:hanging'), '360')
                if qn('w:left') in ind.attrib:
                    del ind.attrib[qn('w:left')]
                
                rPr = lvl.find(qn('w:rPr'))
                if rPr is None:
                    rPr = create_element('w:rPr')
                    lvl.append(rPr)
                rtl = rPr.find(qn('w:rtl'))
                if rtl is None:
                    rtl = create_element('w:rtl')
                    rPr.append(rtl)
                rtl.set(qn('w:val'), '1')

                numFmt = lvl.find(qn('w:numFmt'))
                numFmt_val = numFmt.get(qn('w:val')) if numFmt is not None else ''
                lvlText = lvl.find(qn('w:lvlText'))
                lvlText_val = lvlText.get(qn('w:val')) if lvlText is not None else ''

                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = create_element('w:rFonts')
                    rPr.append(rFonts)
                
                if numFmt_val == 'bullet':
                    if '\uf0b7' in lvlText_val:
                        rFonts.set(qn('w:ascii'), 'Symbol')
                        rFonts.set(qn('w:hAnsi'), 'Symbol')
                        rFonts.set(qn('w:cs'), 'Symbol')
                        rFonts.set(qn('w:hint'), 'default')
                    elif '\uf0a7' in lvlText_val:
                        rFonts.set(qn('w:ascii'), 'Wingdings')
                        rFonts.set(qn('w:hAnsi'), 'Wingdings')
                        rFonts.set(qn('w:cs'), 'Wingdings')
                        rFonts.set(qn('w:hint'), 'default')
                    elif lvlText_val == 'o':
                        rFonts.set(qn('w:ascii'), 'Courier New')
                        rFonts.set(qn('w:hAnsi'), 'Courier New')
                        rFonts.set(qn('w:cs'), 'Courier New')
                        rFonts.set(qn('w:hint'), 'default')
                    else:
                        rFonts.set(qn('w:ascii'), 'Arial')
                        rFonts.set(qn('w:hAnsi'), 'Arial')
                        rFonts.set(qn('w:cs'), 'Arial')
                else:
                    rFonts.set(qn('w:cs'), CFG.font_family_persian)
                    rFonts.set(qn('w:ascii'), CFG.font_family_latin)
                    rFonts.set(qn('w:hAnsi'), CFG.font_family_latin)
            from lxml import etree
            part._blob = etree.tostring(root, encoding='utf-8', xml_declaration=True)

    # Enforce RTL bidi and fonts across ALL document styles in memory
    for style in doc.styles:
        if hasattr(style, '_element'):
            pPr = style._element.get_or_add_pPr()
            bidi = pPr.find(qn('w:bidi'))
            if bidi is None:
                bidi = create_element('w:bidi')
                pPr.append(bidi)
            bidi.set(qn('w:val'), '1')
            
            rPr = style._element.get_or_add_rPr()
            rtl = rPr.find(qn('w:rtl'))
            if rtl is None:
                rtl = create_element('w:rtl')
                rPr.append(rtl)
            rtl.set(qn('w:val'), '1')
            
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = create_element('w:rFonts')
                rPr.append(rFonts)
            rFonts.set(qn('w:cs'), CFG.font_family_persian)
            rFonts.set(qn('w:ascii'), CFG.font_family_latin)
            rFonts.set(qn('w:hAnsi'), CFG.font_family_latin)

    # Format settings.xml for document-wide bidi
    for part in doc.part.package.parts:
        if 'settings' in part.partname:
            from docx.oxml import parse_xml
            root = parse_xml(part.blob)
            bidi = root.find(qn('w:bidi'))
            if bidi is None:
                bidi = create_element('w:bidi')
                root.append(bidi)
            bidi.set(qn('w:val'), '1')
            from lxml import etree
            part._blob = etree.tostring(root, encoding='utf-8', xml_declaration=True)

    # Save to thesis.docx
    doc.save(CFG.output_docx)
    print(f"wrote {CFG.output_docx}")

    # Inject docDefaults RTL bidi directly into word/styles.xml inside zip package
    import zipfile, tempfile
    tmp_fd, tmp_path = tempfile.mkstemp()
    os.close(tmp_fd)
    with zipfile.ZipFile(CFG.output_docx, 'r') as zin:
        with zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                buffer = zin.read(item.filename)
                if item.filename == 'word/document.xml':
                    from lxml import etree
                    root = etree.fromstring(buffer)
                    M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'
                    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    ns = {'m': M_NS, 'w': W_NS}
                    
                    for math in root.xpath('//m:oMath | //m:oMathPara', namespaces=ns):
                        for rPr in math.xpath('.//w:rPr', namespaces=ns):
                            rFonts = rPr.find('{%s}rFonts' % W_NS)
                            if rFonts is None:
                                rFonts = etree.SubElement(rPr, '{%s}rFonts' % W_NS)
                            rFonts.set('{%s}ascii' % W_NS, 'Cambria Math')
                            rFonts.set('{%s}hAnsi' % W_NS, 'Cambria Math')
                            rFonts.set('{%s}cs' % W_NS, 'Cambria Math')
                            
                            rtl = rPr.find('{%s}rtl' % W_NS)
                            if rtl is None:
                                rtl = etree.SubElement(rPr, '{%s}rtl' % W_NS)
                            rtl.set('{%s}val' % W_NS, '0')
                            
                            for tag in ['bCs', 'iCs']:
                                el = rPr.find('{%s}%s' % (W_NS, tag))
                                if el is not None:
                                    rPr.remove(el)
                        
                        for mr in math.xpath('.//m:r', namespaces=ns):
                            rPr = mr.find('{%s}rPr' % W_NS)
                            if rPr is None:
                                rPr = etree.Element('{%s}rPr' % W_NS)
                                mr.insert(0, rPr)
                                rFonts = etree.SubElement(rPr, '{%s}rFonts' % W_NS)
                                rFonts.set('{%s}ascii' % W_NS, 'Cambria Math')
                                rFonts.set('{%s}hAnsi' % W_NS, 'Cambria Math')
                                rFonts.set('{%s}cs' % W_NS, 'Cambria Math')
                                rtl = etree.SubElement(rPr, '{%s}rtl' % W_NS)
                                rtl.set('{%s}val' % W_NS, '0')
                    
                    # Ensure table borders and TableGrid style for all tables
                    for tbl in root.xpath('//w:tbl', namespaces=ns):
                        tblPr = tbl.find('{%s}tblPr' % W_NS)
                        if tblPr is not None:
                            st = tblPr.find('{%s}tblStyle' % W_NS)
                            if st is not None:
                                st.set('{%s}val' % W_NS, 'TableGrid')
                            
                            jc = tblPr.find('{%s}jc' % W_NS)
                            if jc is None:
                                jc = etree.SubElement(tblPr, '{%s}jc' % W_NS)
                            jc.set('{%s}val' % W_NS, 'center')

                            borders = tblPr.find('{%s}tblBorders' % W_NS)
                            if borders is None:
                                borders = etree.SubElement(tblPr, '{%s}tblBorders' % W_NS)
                            for b_name, b_sz in [('top', '8'), ('bottom', '8'), ('left', '6'), ('right', '6'), ('insideH', '4'), ('insideV', '4')]:
                                b = borders.find('{%s}%s' % (W_NS, b_name))
                                if b is None:
                                    b = etree.SubElement(borders, '{%s}%s' % (W_NS, b_name))
                                b.set('{%s}val' % W_NS, 'single')
                                b.set('{%s}sz' % W_NS, b_sz)
                                b.set('{%s}space' % W_NS, '0')
                                b.set('{%s}color' % W_NS, 'auto')
                    
                    buffer = etree.tostring(root, encoding='utf-8', xml_declaration=True)
                elif item.filename == 'word/styles.xml':
                    from lxml import etree
                    root = etree.fromstring(buffer)
                    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    docDef = root.find('{%s}docDefaults' % W_NS)
                    if docDef is not None:
                        pPrDef = docDef.find('{%s}pPrDefault' % W_NS)
                        if pPrDef is not None:
                            pPr = pPrDef.find('{%s}pPr' % W_NS)
                            if pPr is not None:
                                bidi = pPr.find('{%s}bidi' % W_NS)
                                if bidi is None:
                                    bidi = etree.Element('{%s}bidi' % W_NS)
                                    pPr.append(bidi)
                                bidi.set('{%s}val' % W_NS, '1')
                                jc = pPr.find('{%s}jc' % W_NS)
                                if jc is None:
                                    jc = etree.Element('{%s}jc' % W_NS)
                                    pPr.append(jc)
                                jc.set('{%s}val' % W_NS, 'both')
                        rPrDef = docDef.find('{%s}rPrDefault' % W_NS)
                        if rPrDef is not None:
                            rPr = rPrDef.find('{%s}rPr' % W_NS)
                            if rPr is not None:
                                rtl = rPr.find('{%s}rtl' % W_NS)
                                if rtl is None:
                                    rtl = etree.Element('{%s}rtl' % W_NS)
                                    rPr.append(rtl)
                                rtl.set('{%s}val' % W_NS, '1')
                                lang = rPr.find('{%s}lang' % W_NS)
                                if lang is None:
                                    lang = etree.Element('{%s}lang' % W_NS)
                                    rPr.append(lang)
                                lang.set('{%s}bidi' % W_NS, 'fa-IR')
                                rFonts = rPr.find('{%s}rFonts' % W_NS)
                                if rFonts is None:
                                    rFonts = etree.Element('{%s}rFonts' % W_NS)
                                    rPr.append(rFonts)
                                rFonts.set('{%s}cs' % W_NS, CFG.font_family_persian)
                                rFonts.set('{%s}ascii' % W_NS, CFG.font_family_latin)
                                rFonts.set('{%s}hAnsi' % W_NS, CFG.font_family_latin)
                    
                    # Ensure Normal, BodyText, and all paragraph styles have bidi, rtl, and fa-IR
                    for s in root.iter('{%s}style' % W_NS):
                        s_type = s.get('{%s}type' % W_NS)
                        s_id = s.get('{%s}styleId' % W_NS, '')
                        if s_type == 'paragraph':
                            pPr = s.find('{%s}pPr' % W_NS)
                            if pPr is None:
                                pPr = etree.SubElement(s, '{%s}pPr' % W_NS)
                            bidi = pPr.find('{%s}bidi' % W_NS)
                            if bidi is None:
                                bidi = etree.Element('{%s}bidi' % W_NS)
                                pPr.append(bidi)
                            bidi.set('{%s}val' % W_NS, '1')
                            if s_id in ['Normal', 'BodyText', 'Body Text']:
                                jc = pPr.find('{%s}jc' % W_NS)
                                if jc is None:
                                    jc = etree.Element('{%s}jc' % W_NS)
                                    pPr.append(jc)
                                jc.set('{%s}val' % W_NS, 'both')
                            elif 'heading' in s_id.lower():
                                jc = pPr.find('{%s}jc' % W_NS)
                                if jc is None:
                                    jc = etree.Element('{%s}jc' % W_NS)
                                    pPr.append(jc)
                                jc.set('{%s}val' % W_NS, 'right')
                            
                            rPr = s.find('{%s}rPr' % W_NS)
                            if rPr is None:
                                rPr = etree.SubElement(s, '{%s}rPr' % W_NS)
                            rtl = rPr.find('{%s}rtl' % W_NS)
                            if rtl is None:
                                rtl = etree.Element('{%s}rtl' % W_NS)
                                rPr.append(rtl)
                            rtl.set('{%s}val' % W_NS, '1')
                            lang = rPr.find('{%s}lang' % W_NS)
                            if lang is None:
                                lang = etree.Element('{%s}lang' % W_NS)
                                rPr.append(lang)
                            lang.set('{%s}bidi' % W_NS, 'fa-IR')
                            
                            if 'heading' in s_id.lower():
                                color = rPr.find('{%s}color' % W_NS)
                                if color is None:
                                    color = etree.Element('{%s}color' % W_NS)
                                    rPr.append(color)
                                color.set('{%s}val' % W_NS, '000000')
                                if '{%s}themeColor' % W_NS in color.attrib:
                                    del color.attrib['{%s}themeColor' % W_NS]
                                if '{%s}themeShade' % W_NS in color.attrib:
                                    del color.attrib['{%s}themeShade' % W_NS]
                                
                                i_el = rPr.find('{%s}i' % W_NS)
                                if i_el is not None:
                                    rPr.remove(i_el)
                                iCs_el = rPr.find('{%s}iCs' % W_NS)
                                if iCs_el is not None:
                                    rPr.remove(iCs_el)
                                
                                b_el = rPr.find('{%s}b' % W_NS)
                                if b_el is None:
                                    b_el = etree.Element('{%s}b' % W_NS)
                                    rPr.append(b_el)
                                bCs_el = rPr.find('{%s}bCs' % W_NS)
                                if bCs_el is None:
                                    bCs_el = etree.Element('{%s}bCs' % W_NS)
                                    rPr.append(bCs_el)

                        elif s_type == 'character' and 'heading' in s_id.lower():
                            rPr = s.find('{%s}rPr' % W_NS)
                            if rPr is not None:
                                color = rPr.find('{%s}color' % W_NS)
                                if color is None:
                                    color = etree.Element('{%s}color' % W_NS)
                                    rPr.append(color)
                                color.set('{%s}val' % W_NS, '000000')
                                if '{%s}themeColor' % W_NS in color.attrib:
                                    del color.attrib['{%s}themeColor' % W_NS]
                                if '{%s}themeShade' % W_NS in color.attrib:
                                    del color.attrib['{%s}themeShade' % W_NS]
                                i_el = rPr.find('{%s}i' % W_NS)
                                if i_el is not None:
                                    rPr.remove(i_el)
                                iCs_el = rPr.find('{%s}iCs' % W_NS)
                                if iCs_el is not None:
                                    rPr.remove(iCs_el)
                    
                    # Ensure FootnoteReference character style in styles.xml
                    fn_ref_style = root.find('.//{%s}style[@{%s}styleId="FootnoteReference"]' % (W_NS, W_NS))
                    if fn_ref_style is None:
                        fn_ref_style = etree.SubElement(root, '{%s}style' % W_NS)
                        fn_ref_style.set('{%s}type' % W_NS, 'character')
                        fn_ref_style.set('{%s}styleId' % W_NS, 'FootnoteReference')
                        name_el = etree.SubElement(fn_ref_style, '{%s}name' % W_NS)
                        name_el.set('{%s}val' % W_NS, 'Footnote Reference')
                    rPr_fn = fn_ref_style.find('{%s}rPr' % W_NS)
                    if rPr_fn is None:
                        rPr_fn = etree.SubElement(fn_ref_style, '{%s}rPr' % W_NS)
                    vert = rPr_fn.find('{%s}vertAlign' % W_NS)
                    if vert is None:
                        vert = etree.SubElement(rPr_fn, '{%s}vertAlign' % W_NS)
                    vert.set('{%s}val' % W_NS, 'superscript')
                    
                    buffer = etree.tostring(root, encoding='utf-8', xml_declaration=True)
                elif item.filename == 'word/numbering.xml':
                    from lxml import etree
                    root = etree.fromstring(buffer)
                    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    ns = {'w': W_NS}
                    for lvl in root.xpath('//w:lvl', namespaces=ns):
                        ilvl = int(lvl.get('{%s}ilvl' % W_NS, '0'))
                        
                        # 1. Level justification right
                        lvlJc = lvl.find('{%s}lvlJc' % W_NS)
                        if lvlJc is None:
                            lvlJc = etree.SubElement(lvl, '{%s}lvlJc' % W_NS)
                        lvlJc.set('{%s}val' % W_NS, 'right')
                        
                        # 2. pPr with RTL bidi and right indent
                        pPr = lvl.find('{%s}pPr' % W_NS)
                        if pPr is None:
                            pPr = etree.SubElement(lvl, '{%s}pPr' % W_NS)
                        bidi = pPr.find('{%s}bidi' % W_NS)
                        if bidi is None:
                            bidi = etree.SubElement(pPr, '{%s}bidi' % W_NS)
                        bidi.set('{%s}val' % W_NS, '1')
                        
                        ind = pPr.find('{%s}ind' % W_NS)
                        if ind is None:
                            ind = etree.SubElement(pPr, '{%s}ind' % W_NS)
                        right_val = 720 * (ilvl + 1)
                        ind.set('{%s}right' % W_NS, str(right_val))
                        ind.set('{%s}hanging' % W_NS, '360')
                        if '{%s}left' % W_NS in ind.attrib:
                            del ind.attrib['{%s}left' % W_NS]
                        
                        # 3. rPr with RTL
                        rPr = lvl.find('{%s}rPr' % W_NS)
                        if rPr is None:
                            rPr = etree.SubElement(lvl, '{%s}rPr' % W_NS)
                        rtl = rPr.find('{%s}rtl' % W_NS)
                        if rtl is None:
                            rtl = etree.SubElement(rPr, '{%s}rtl' % W_NS)
                        rtl.set('{%s}val' % W_NS, '1')
                        lang = rPr.find('{%s}lang' % W_NS)
                        if lang is None:
                            lang = etree.Element('{%s}lang' % W_NS)
                            rPr.append(lang)
                        lang.set('{%s}bidi' % W_NS, 'fa-IR')
                        
                        numFmt = lvl.find('{%s}numFmt' % W_NS)
                        numFmt_val = numFmt.get('{%s}val' % W_NS) if numFmt is not None else ''
                        lvlText = lvl.find('{%s}lvlText' % W_NS)
                        lvlText_val = lvlText.get('{%s}val' % W_NS) if lvlText is not None else ''

                        rFonts = rPr.find('{%s}rFonts' % W_NS)
                        if rFonts is None:
                            rFonts = etree.SubElement(rPr, '{%s}rFonts' % W_NS)
                        
                        if numFmt_val == 'bullet':
                            if '\uf0b7' in lvlText_val:
                                rFonts.set('{%s}ascii' % W_NS, 'Symbol')
                                rFonts.set('{%s}hAnsi' % W_NS, 'Symbol')
                                rFonts.set('{%s}cs' % W_NS, 'Symbol')
                                rFonts.set('{%s}hint' % W_NS, 'default')
                            elif '\uf0a7' in lvlText_val:
                                rFonts.set('{%s}ascii' % W_NS, 'Wingdings')
                                rFonts.set('{%s}hAnsi' % W_NS, 'Wingdings')
                                rFonts.set('{%s}cs' % W_NS, 'Wingdings')
                                rFonts.set('{%s}hint' % W_NS, 'default')
                            elif lvlText_val == 'o':
                                rFonts.set('{%s}ascii' % W_NS, 'Courier New')
                                rFonts.set('{%s}hAnsi' % W_NS, 'Courier New')
                                rFonts.set('{%s}cs' % W_NS, 'Courier New')
                                rFonts.set('{%s}hint' % W_NS, 'default')
                            else:
                                rFonts.set('{%s}ascii' % W_NS, 'Arial')
                                rFonts.set('{%s}hAnsi' % W_NS, 'Arial')
                                rFonts.set('{%s}cs' % W_NS, 'Arial')
                        else:
                            rFonts.set('{%s}cs' % W_NS, CFG.font_family_persian)
                            rFonts.set('{%s}ascii' % W_NS, CFG.font_family_latin)
                            rFonts.set('{%s}hAnsi' % W_NS, CFG.font_family_latin)
                    
                    buffer = etree.tostring(root, encoding='utf-8', xml_declaration=True)
                elif item.filename == 'word/settings.xml':
                    from lxml import etree
                    root = etree.fromstring(buffer)
                    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                    
                    # Ensure bidi tag in settings
                    bidi = root.find('{%s}bidi' % W_NS)
                    if bidi is None:
                        bidi = etree.Element('{%s}bidi' % W_NS)
                        root.append(bidi)
                    bidi.set('{%s}val' % W_NS, '1')
                    
                    # Ensure themeFontLang has bidi="fa-IR"
                    tfl = root.find('{%s}themeFontLang' % W_NS)
                    if tfl is None:
                        tfl = etree.SubElement(root, '{%s}themeFontLang' % W_NS)
                    tfl.set('{%s}bidi' % W_NS, 'fa-IR')
                    tfl.set('{%s}val' % W_NS, 'en-US')
                    
                    uf = root.find('{%s}updateFields' % W_NS)
                    if uf is None:
                        uf = etree.Element('{%s}updateFields' % W_NS)
                        root.append(uf)
                    uf.set('{%s}val' % W_NS, 'true')
                    buffer = etree.tostring(root, encoding='utf-8', xml_declaration=True)
                zout.writestr(item, buffer)
    shutil.move(tmp_path, CFG.output_docx)

    # Inject dual-section page numbering via direct zip manipulation
    _inject_page_numbering(CFG.output_docx, CFG.font_family_persian, CFG.font_family_latin)

if __name__ == "__main__":
    postprocess_doc()
