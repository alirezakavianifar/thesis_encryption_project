"""Update main_updated.docx with numerical results from code.ipynb / results.pkl."""
from docx import Document
from pathlib import Path

DOCX_PATH = Path(r"e:\projects\thesis_project_v2\main_updated.docx")
OUT_PATH = DOCX_PATH  # overwrite in place when not locked

# Table 4-2: original entropy values (encrypted values already correct)
ENTROPY_REPLACEMENTS = {
    "6.2138 → 7.9993": "7.2542 → 7.9993",
    "6.7990 → 7.9993": "7.1491 → 7.9993",
    "6.7178 → 7.9993": "5.8825 → 7.9993",
    "7.7522 → 7.9993": "1.0000 → 7.9993",
    "7.4744 → 7.9993": "7.5754 → 7.9993",
    "7.7067 → 7.9994": "7.2855 → 7.9994",
    "7.0583 → 7.9994": "7.9931 → 7.9994",
    "7.4963 → 7.9992": "7.9931 → 7.9992",
    "7.3388 → 7.9993": "7.6718 → 7.9993",
    "6.9207 → 7.9993": "7.6075 → 7.9993",
    "7.4136 → 7.9993": "7.6398 → 7.9993",
    "7.2104 → 7.9993": "6.3472 → 7.9993",
}

TEXT_REPLACEMENTS = [
    ("از ۶.۵۷ به ۷.۹۹۹۳ بیت", "از ۶.۷۶ به ۷.۹۹۹۳ بیت"),
    ("آنتروپی میانگین ~۶.۵۷ بیت", "آنتروپی میانگین ~۶.۷۶ بیت"),
    ("مقادیر تصویر Airplane (۲۱.۲%)", "مقادیر تصویر Airplane (۲۶.۳۴٪)"),
    ("Baboon ۳۱.۳٪، Tree: ۳۱.۱٪", "Baboon ۳۰.۲۰٪، Tree: ۳۱.۳۱٪"),
    (
        "NPCR=99.46%برای تصویر Baboon در محدوده عملکرد روش‌های مرجع قرار دارد.",
        "NPCR تا ۹۹.۴۸٪ برای تصویر Peppers و میانگین ۹۴.۹۲٪ در محدوده عملکرد روش‌های مرجع قرار دارد.",
    ),
    ("30/19", "30.19"),
    ("94/91", "94.92"),
    ("7/9993", "7.9993"),
]


def replace_in_paragraph(paragraph, replacements):
    full = paragraph.text
    if not full:
        return False
    new = full
    for old, new_val in replacements:
        if old in new:
            new = new.replace(old, new_val)
    if new == full:
        return False
    # Preserve first run formatting where possible
    if paragraph.runs:
        paragraph.runs[0].text = new
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new
    return True


def main():
    doc = Document(DOCX_PATH)
    counts = {"entropy": 0, "text": 0}

    all_replacements = list(ENTROPY_REPLACEMENTS.items()) + TEXT_REPLACEMENTS

    for para in doc.paragraphs:
        for old, new in ENTROPY_REPLACEMENTS.items():
            if old in para.text:
                replace_in_paragraph(para, [(old, new)])
                counts["entropy"] += 1
                break
        else:
            for old, new in TEXT_REPLACEMENTS:
                if old in para.text:
                    replace_in_paragraph(para, [(old, new)])
                    counts["text"] += 1
                    break

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in ENTROPY_REPLACEMENTS.items():
                        if old in para.text:
                            replace_in_paragraph(para, [(old, new)])
                            counts["entropy"] += 1
                            break
                    else:
                        for old, new in TEXT_REPLACEMENTS:
                            if old in para.text:
                                replace_in_paragraph(para, [(old, new)])
                                counts["text"] += 1
                                break

    try:
        doc.save(OUT_PATH)
        saved = OUT_PATH
    except PermissionError:
        saved = DOCX_PATH.with_name("main_updated_corrected.docx")
        doc.save(saved)
        print("WARNING: original file locked; saved copy instead.")
    print("Updated:", saved)
    print("Counts:", counts)


if __name__ == "__main__":
    main()
