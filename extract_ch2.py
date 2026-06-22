"""Find body content for chapter 2 section 2-6."""
from docx import Document
from pathlib import Path

doc = Document(Path(r"e:\projects\thesis_project_v2\main_updated.docx"))
out = Path(r"e:\projects\thesis_project_v2\ch2_body.txt")

for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    style = p.style.name if p.style else ""
    if 480 <= i <= 540 or "پیشینه" in t or "2-6" in t or "مطالعات انجام" in t:
        if t:
            out.write_text("", encoding="utf-8") if i == 480 else None

lines = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    style = p.style.name if p.style else ""
    if i >= 470 and i <= 560:
        if t:
            lines.append(f"[{i}|{style}] {t}")

# references section
lines.append("\n=== REFS (last 40 paras) ===\n")
for i, p in enumerate(doc.paragraphs):
    if i >= len(doc.paragraphs) - 50:
        t = p.text.strip()
        if t and ("[" in t[:5] or "مراجع" in t):
            lines.append(f"[{i}] {t[:200]}")

with open(out, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))
print("done", len(lines))
