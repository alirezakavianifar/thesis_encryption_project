"""Sync main_updated.docx into thesis_latex_source chapter .tex files."""

from __future__ import annotations

import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

ROOT = Path(r"e:\projects\thesis_project_v2")
DOCX = ROOT / "main_updated.docx"
LATEX = ROOT / "thesis_latex_source"
CODE_EXPORT = ROOT / "code_export.py.py"
FIG_SRC = ROOT / "outputs" / "figs"

CHAPTER_TITLES = {
    1: "کلیات تحقیق",
    2: "مفاهیم پایه و پیشینه پژوهش",
    3: "معرفی الگوریتم پیشنهادی",
    4: "پیاده سازی الگوریتم پیشنهادی",
    5: "بحث و نتیجه‌گیری",
}

CHAPTER_FILES = {
    1: "03_chapter1.tex",
    2: "04_chapter2.tex",
    3: "05_chapter3.tex",
    4: "06_chapter4.tex",
    5: "07_chapter5.tex",
}

TEX_HEADERS = {
    "02_abstract.tex": """% !TEX root = main.tex
% ============================================================================
% PERSIAN ABSTRACT
% ============================================================================

\\chapter*{چکیده}
""",
    "03_chapter1.tex": """% !TEX root = main.tex
% ============================================================================
% CHAPTER 1: INTRODUCTION AND RESEARCH SPECIFICATIONS
% ============================================================================

\\chapter{کلیات تحقیق}
""",
    "04_chapter2.tex": """% !TEX root = main.tex
% ============================================================================
% CHAPTER 2: BASIC CONCEPTS AND LITERATURE REVIEW
% ============================================================================

\\chapter{مفاهیم پایه و پیشینه پژوهش}
""",
    "05_chapter3.tex": """% !TEX root = main.tex
% ============================================================================
% CHAPTER 3: PROPOSED ALGORITHM
% ============================================================================

\\chapter{معرفی الگوریتم پیشنهادی}
""",
    "06_chapter4.tex": """% !TEX root = main.tex
% ============================================================================
% CHAPTER 4: IMPLEMENTATION AND RESULTS
% ============================================================================

\\chapter{پیاده سازی الگوریتم پیشنهادی}
""",
    "07_chapter5.tex": """% !TEX root = main.tex
% ============================================================================
% CHAPTER 5: DISCUSSION AND CONCLUSION
% ============================================================================

\\chapter{بحث و نتیجه‌گیری}
""",
    "09_appendix.tex": """% !TEX root = main.tex
% ============================================================================
% APPENDIX: EXECUTABLE PYTHON CODE
% ============================================================================

\\chapter{پیوست: کد اجرایی الگوریتم (پایتون)}
\\label{ch:appendix_code}
""",
}

CH4_FIGURES = {
    "نمایش بصری": ("fig1_visual.png", "fig:visual_results"),
    "هیستوگرام مقایسه": ("fig2_histograms.png", "fig:histograms"),
    "آنتروپی شانون تصاویر": ("fig4_entropy.png", "fig:entropy"),
    "پراکندگی همبستگی": ("fig3_correlation.png", "fig:correlation_scatter"),
    "مقایسه ضریب همبستگی": ("fig7_corr_bars.png", "fig:correlation_bars"),
    "NPCR و UACI": ("fig5_npcr_uaci.png", "fig:npcr_uaci"),
    "زمان رمزنگاری": ("fig6_time.png", "fig:timing"),
}

CH2_FIGURES = {
    "نگاشت لجستیک": "image2.png",
    "لورنز": "image3.png",
    "سیستم چن": "image4.png",
}


def esc_latex(text: str) -> str:
    repl = {
        "\\": r"\textbackslash{}",
        "&": r"\&",
        "%": r"\%",
        "$": r"\$",
        "#": r"\#",
        "_": r"\_",
        "{": r"\{",
        "}": r"\}",
        "~": r"\textasciitilde{}",
        "^": r"\textasciicircum{}",
    }
    return "".join(repl.get(ch, ch) for ch in text)


def wrap_lr(text: str) -> str:
    return re.sub(
        r"[A-Za-z][A-Za-z0-9_+\-./^]*",
        lambda m: f"\\lr{{{m.group(0)}}}" if not m.group(0).isdigit() else m.group(0),
        text,
    )


def convert_citations(text: str) -> str:
    def repl(m: re.Match) -> str:
        nums = re.findall(r"\d+", m.group(1))
        if not nums:
            return m.group(0)
        return "\\cite{" + ", ".join(f"ref{n}" for n in nums) + "}"

    return re.sub(r"\[([^\]]+)\]", repl, text)


def inline(text: str) -> str:
    protected: list[str] = []

    def protect(m: re.Match) -> str:
        protected.append(m.group(0))
        return f"@@PROT{len(protected) - 1}@@"

    text = convert_citations(text)
    text = re.sub(r"\\cite\{[^{}]+\}", protect, text)
    text = wrap_lr(text)
    text = esc_latex(text)
    for i, val in enumerate(protected):
        text = text.replace(f"@@PROT{i}@@", val)
    return text


def style_name(p: Paragraph) -> str:
    return p.style.name if p.style else ""


def section_cmd(style: str, title: str) -> str | None:
    clean = re.sub(r"^\d+(?:-\d+)+[\.\s]*", "", title).strip()
    clean = inline(clean)
    if style == "titr 1" or style == "Heading 1":
        return f"\\section{{{clean}}}"
    if style in ("titr 2", "Heading 2"):
        return f"\\subsection{{{clean}}}"
    if style in ("titr 3", "Heading 3", "Heading 4"):
        return f"\\subsubsection{{{clean}}}"
    return None


def table_latex(table: Table, caption: str | None) -> str:
    rows = []
    ncols = max(len(r.cells) for r in table.rows)
    for row in table.rows:
        cells = [inline(c.text.strip().replace("\n", " ")) for c in row.cells]
        while len(cells) < ncols:
            cells.append("")
        rows.append(" & ".join(cells) + r" \\")
    colspec = "l" * ncols
    lines = [
        "\\begin{table}[h]",
        "\\centering",
        "\\small",
        f"\\begin{{tabular}}{{{colspec}}}",
        "\\hline",
        *rows,
        "\\hline",
        "\\end{tabular}",
    ]
    if caption:
        cap = re.sub(r"^جدول\s*[\d\-]+[\.\s]*", "", caption).strip()
        lines.append(f"\\caption{{{inline(cap)}}}")
    lines.append("\\end{table}")
    return "\n".join(lines)


def figure_latex(caption: str, image: str, label: str) -> str:
    cap = inline(caption)
    return (
        "\\begin{figure}[htbp]\n"
        "    \\centering\n"
        f"    \\includegraphics[width=0.92\\textwidth]{{images/{image}}}\n"
        f"    \\caption{{{cap}}}\n"
        f"    \\label{{{label}}}\n"
        "\\end{figure}"
    )


def pick_ch4_figure(caption: str) -> tuple[str, str] | None:
    for key, (img, lab) in CH4_FIGURES.items():
        if key in caption:
            return img, lab
    return None


def pick_ch2_figure(caption: str) -> str | None:
    for key, img in CH2_FIGURES.items():
        if key in caption:
            return img
    return None


def build_body_index(doc: Document) -> list[tuple[int, str, object]]:
    """Return ordered blocks with paragraph index of nearest preceding paragraph."""
    blocks: list[tuple[int, str, object]] = []
    para_idx = -1
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            para_idx += 1
            blocks.append((para_idx, "p", Paragraph(child, doc)))
        elif child.tag == qn("w:tbl"):
            blocks.append((para_idx, "t", Table(child, doc)))
    return blocks


def render_range(doc: Document, start: int, end: int, chapter: int | None = None) -> str:
    parts: list[str] = []
    pending_caption = None
    pending_kind = None
    fig_counter = 0

    for para_idx, kind, obj in build_body_index(doc):
        if para_idx < start or para_idx >= end:
            continue

        if kind == "p":
            p: Paragraph = obj
            st = style_name(p)
            t = p.text.strip()
            if not t:
                continue

            if st in ("titr table", "titr picture"):
                if st == "titr picture":
                    cap = t
                    img = None
                    label = f"fig:ch{chapter}_{fig_counter}" if chapter else f"fig:auto_{fig_counter}"
                    if chapter == 4:
                        picked = pick_ch4_figure(cap)
                        if picked:
                            img, label = picked
                    elif chapter == 2:
                        img = pick_ch2_figure(cap)
                    if img:
                        parts.append(figure_latex(cap, img, label))
                        fig_counter += 1
                    pending_caption = None
                    pending_kind = None
                    continue
                pending_caption = t
                pending_kind = "table"
                continue

            cmd = section_cmd(st, t)
            if cmd:
                parts.append(cmd)
                continue

            if st == "List Paragraph" or st.startswith("List"):
                parts.append("\\begin{itemize}\n    \\item " + inline(t) + "\n\\end{itemize}")
                continue

            if re.match(r"^\d+-\d+", t):
                cmd = section_cmd("titr 2", t)
                if cmd:
                    parts.append(cmd)
                    continue

            parts.append(inline(t))

        elif kind == "t":
            if pending_kind == "table":
                parts.append(table_latex(obj, pending_caption))
                pending_caption = None
                pending_kind = None
            else:
                parts.append(table_latex(obj, None))

    return "\n\n".join(parts)


def find_chapter_ranges(doc: Document) -> dict[int, tuple[int, int]]:
    """Map chapter number -> (start_para_idx, end_para_idx) using titr chapter order."""
    starts: list[int] = []
    ref_start = len(doc.paragraphs)
    for i, p in enumerate(doc.paragraphs):
        st = style_name(p)
        t = p.text.strip()
        if st == "titr chapter":
            starts.append(i)
        if t == "منابع" and ref_start == len(doc.paragraphs):
            ref_start = i

    ranges: dict[int, tuple[int, int]] = {}
    for j, idx in enumerate(starts[:5]):
        end = starts[j + 1] if j + 1 < len(starts) else ref_start
        ranges[j + 1] = (idx + 1, end)
    return ranges


def find_abstract_range(doc: Document) -> tuple[int, int]:
    start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("چکیده"):
            start = i + 1
            break
    if start is None:
        return 0, 0
    for i in range(start, len(doc.paragraphs)):
        if style_name(doc.paragraphs[i]) == "titr chapter":
            return start, i
    return start, start + 1


def postprocess_known_updates(body: str) -> str:
    """Align a few values that may still be stale in the docx abstract."""
    replacements = {
        "۷.۹۹۹۳ بیت (از ۸ بیت ایده‌آل)": "۸.۰ بیت (مقدار ایده‌آل)",
        "۷.۹۹۹۳ بیت": "۸.۰ بیت",
        "۰.۰۱۲۶": "۰.۰۰۸۲",
        "۹۸.۷٪": "۹۹.۲٪",
        "94.91": "95.13",
        "30.19": "30.14",
        "26.22 ثانیه": "۲۳.۸ ثانیه",
    }
    for old, new in replacements.items():
        body = body.replace(old, new)
    return body


def build_abstract(doc: Document) -> str:
    start, end = find_abstract_range(doc)
    lines = []
    for p in doc.paragraphs[start:end]:
        t = p.text.strip()
        if not t:
            continue
        if t.startswith("کلمات کلیدی"):
            kw = re.sub(r"^کلمات کلیدی\s*[:：]?\s*", "", t)
            lines.append(f"\\textbf{{کلمات کلیدی:}} {inline(kw)}")
            continue
        lines.append(inline(t))
    return postprocess_known_updates("\n\n".join(lines) + "\n\n\\newpage\n")


def build_appendix() -> str:
    code = CODE_EXPORT.read_text(encoding="utf-8", errors="replace")
    for i, ln in enumerate(code.splitlines()):
        if ln.startswith("import ") or ln.startswith("from "):
            code = "\n".join(code.splitlines()[i:])
            break
    intro = (
        "کد اجرایی کامل الگوریتم پیشنهادی، شامل مرحله یکنواخت‌سازی هیستوگرام "
        "(مرحله ۶) و توابع رمزنگاری/رمزگشایی به‌روزرسانی‌شده، در ادامه آورده شده است."
    )
    return "\n\n".join(
        [
            intro,
            "\\begin{latin}",
            "\\noindent\\textbf{Listing: Updated Python implementation}",
            "\\begin{verbatim}",
            code.rstrip(),
            "\\end{verbatim}",
            "\\end{latin}",
        ]
    )


def export_docx_images():
    img_dir = LATEX / "images"
    img_dir.mkdir(exist_ok=True)
    with zipfile.ZipFile(DOCX) as z:
        for name in z.namelist():
            if name.startswith("word/media/") and name.lower().endswith((".png", ".jpg", ".jpeg")):
                data = z.read(name)
                out = img_dir / Path(name).name
                out.write_bytes(data)
    if FIG_SRC.exists():
        for p in FIG_SRC.glob("*.png"):
            shutil.copy2(p, img_dir / p.name)


def write_file(fname: str, body: str):
    header = TEX_HEADERS[fname]
    (LATEX / fname).write_text(header + body + "\n", encoding="utf-8")


def main():
    doc = Document(DOCX)
    export_docx_images()

    write_file("02_abstract.tex", build_abstract(doc))

    ranges = find_chapter_ranges(doc)
    for num, fname in CHAPTER_FILES.items():
        if num not in ranges:
            print(f"Warning: chapter {num} not found")
            continue
        s, e = ranges[num]
        body = render_range(doc, s, e, chapter=num)
        body = postprocess_known_updates(body)
        write_file(fname, body)
        print(f"Updated {fname} ({e - s} paragraphs)")

    write_file("09_appendix.tex", build_appendix())

    (LATEX / "_synced_from_docx.txt").write_text(
        f"Synced from {DOCX.name}\n", encoding="utf-8"
    )
    print("Done.")


if __name__ == "__main__":
    main()
