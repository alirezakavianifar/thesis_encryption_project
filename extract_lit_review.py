"""Extract section 2-6 literature review and table 4-7."""
from docx import Document
from pathlib import Path

doc = Document(Path(r"e:\projects\thesis_project_v2\main_updated.docx"))
out = Path(r"e:\projects\thesis_project_v2\lit_review_content.txt")

lines = []
in_section = False
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if "2-6" in t and "پیشینه" in t:
        in_section = True
    if in_section and t.startswith("فصل3"):
        break
    if in_section and t:
        lines.append(f"[{i}] {t}")

lines.append("\n\n=== TABLES ===\n")
for ti, table in enumerate(doc.tables):
  rows = []
  for row in table.rows:
    cells = [c.text.strip().replace("\n", " | ") for c in row.cells]
    rows.append(" || ".join(cells))
  text = "\n".join(rows)
  if any(k in text for k in ["مرجع", "NPCR", "آنتروپی", "Chen", "Lorenz", "پیشینه", "مقایسه"]):
    lines.append(f"\n--- Table {ti} ---\n{text}")

with open(out, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))
print(f"Wrote {out}, {len(lines)} lines")
