"""Add comparative literature table and analysis to section 2-6 in main_updated.docx."""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from pathlib import Path

DOCX = Path(r"e:\projects\thesis_project_v2\main_updated.docx")


def insert_paragraph_after(paragraph, text="", style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        try:
            new_para.style = style
        except Exception:
            pass
    return new_para


def add_table_after(paragraph, rows, cols):
    doc = paragraph.part.document
    table = doc.add_table(rows=rows, cols=cols)
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    paragraph._p.addnext(tbl)
    return table


def set_cell(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold


def main():
    doc = Document(DOCX)

    anchor_idx = None
    for i, p in enumerate(doc.paragraphs):
        if "در سال 2024" in p.text and "[13]" in p.text:
            anchor_idx = i
            break
    if anchor_idx is None:
        for i, p in enumerate(doc.paragraphs):
            if "[12]" in p.text and "فوق" in p.text:
                anchor_idx = i
                break
    if anchor_idx is None:
        raise RuntimeError("Could not find literature review anchor paragraph")

    anchor = doc.paragraphs[anchor_idx]

    for p in doc.paragraphs:
        if "جدول ۲-۳" in p.text and "مقایسه تطبیقی" in p.text:
            print("Comparison section already present; skipping.")
            return

    h = insert_paragraph_after(anchor, "2-6-1. مقایسه تطبیقی پژوهش‌های مرجع")

    paras = [
        "با توجه به مرور پژوهش‌های فوق، روش‌های رمزنگاری تصویر رنگی مبتنی بر آشوب را می‌توان از منظر سیستم(های) آشوبی به‌کاررفته، سازوکار جایگشت و انتشار، نوآوری الگوریتمی و معیارهای امنیتی گزارش‌شده مقایسه کرد. جدول ۲-۳ خلاصه‌ای تطبیقی از مهم‌ترین مطالعات مرتبط با موضوع پژوهش حاضر ارائه می‌دهد. مقادیر عددی آنتروپی، NPCR و UACI تنها در صورتی درج شده‌اند که در مقاله مرجع یا در ارزیابی مشابه گزارش شده باشند؛ در غیر این صورت با «—» مشخص شده است.",
        "همان‌طور که جدول نشان می‌دهد، بسیاری از روش‌های پیشین از یک یا دو سیستم آشوبی ثابت (مانند Logistic، Tent یا Lorenz) بهره می‌برند و انتخاب سیستم رمزنگاری در سراسر تصویر یکسان است. در مقابل، روش پیشنهادی حاضر با ترکیب سیستم سه‌بعدی Chen برای جایگشت و نه سیستم آشوبی نمایی (ECM) با انتخاب پویا در سطح هر پیکسل، لایه‌ای اضافی از پیچیدگی و غیرقابل‌پیش‌بینی بودن ایجاد می‌کند. از نظر آنتروپی شانون، الگوریتم پیشنهادی به مقدار ایده‌آل ۸.۰ بیت در تمام کانال‌های رنگی دست یافته که نسبت به اکثر روش‌های مرجع (حدود ۷.۹۹۷) برتری جزئی اما معنادار دارد و نشان‌دهنده توزیع کاملاً یکنواخت مقادیر پیکسل پس از رمزنگاری است.",
        "در معیار NPCR (حساسیت به کلید)، روش‌های مبتنی بر Lorenz، Chen+Logistic و فوق‌آشوب مقادیر نزدیک به ۹۹.۶٪ گزارش کرده‌اند. میانگین NPCR روش پیشنهادی (۹۵.۱۳٪) در محدوده قابل قبول قرار دارد، اما برای تصویر Baboon مقدار پایین‌تری مشاهده می‌شود که می‌تواند به حساسیت بیشتر سیستم‌های پیوسته Chen نسبت به شرایط اولیه در تصاویر با بافت پیچیده نسبت داده شود. از این رو، اگرچه الگوریتم پیشنهادی از نظر آنتروپی و حذف همبستگی مکانی عملکرد برجسته‌ای دارد، بهبود یکنواختی NPCR در تمام تصاویر آزمایشی از جمله موارد پیشنهادی برای تحقیقات آینده است.",
        "در معیار UACI، روش خدادادی و زندوکیلی [۷]، وانگ و لی [۸] و الگوریتم فوق‌آشوب [۱۲] به مقادیر نزدیک به ۳۳.۴۶٪ (مقدار ایده‌آل) دست یافته‌اند. میانگین UACI روش پیشنهادی (۳۰.۱۴٪) اندکی پایین‌تر است؛ با این حال برای تصویر Peppers مقدار ۳۲.۸۸٪ به ایده‌آل نزدیک است. در مقایسه با روش لورنز [۶] که UACI مشابه (۳۰.۳۵٪) گزارش کرده، عملکرد روش پیشنهادی قابل رقابت است.",
        "از منظر کارایی، الگوریتم فوق‌آشوب [۱۲] با زمان رمزنگاری ۰.۰۲۳ ثانیه سرعت بسیار بالایی دارد؛ این برتری عمدتاً ناشی از برداری‌سازی و کاهش تکرارهای آشوبی است. پیاده‌سازی فعلی روش پیشنهادی در Python با حلقه وابسته به انتخاب پویای ECM برای هر پیکسل، زمان رمزنگاری حدود ۲۳–۲۵ ثانیه برای تصویر ۵۱۲×۵۱۲ ایجاد می‌کند که برای اثبات مفهوم مناسب است اما در کاربرد بلادرنگ نیازمند بهینه‌سازی خواهد بود.",
        "به‌طور خلاصه، نوآوری اصلی پژوهش حاضر نسبت به ادبیات، «انتخاب پویا و مستقل یکی از نه سیستم نمایی برای هر پیکسل» در کنار «جایگشت مبتنی بر سیستم پیوسته Chen» و «دستیابی به آنتروپی و هیستوگرام کاملاً یکنواخت» است. این ترکیب، روش پیشنهادی را از روش‌های تک‌آشوبی و حتی روش‌های دوگانه (Chen+Logistic) متمایز می‌سازد؛ هرچند در NPCR و UACI هنوز فاصله‌ای با بهترین روش‌های مرجع وجود دارد که در فصل‌های چهارم و پنجم با جزئیات بیشتر مورد بحث قرار گرفته است.",
    ]

    last = h
    for text in paras:
        last = insert_paragraph_after(last, text)

    cap = insert_paragraph_after(last, "جدول ۲-۳. مقایسه تطبیقی روش‌های رمزنگاری تصویر رنگی مبتنی بر آشوب")
    try:
        cap.style = "titr table"
    except Exception:
        pass

    headers = [
        "مرجع",
        "سیستم(های) آشوبی",
        "جایگشت",
        "انتشار",
        "ویژگی متمایز",
        "آنتروپی",
        "NPCR (%)",
        "UACI (%)",
    ]
    rows_data = [
        ["ژانگ و لیو [۳]", "Henon + Tent", "بله", "بله", "ترکیب دو نگاشت گسسته", "—", "—", "—"],
        ["هو و ژو [۴]", "۹ ECM", "—", "—", "پایداری آشوب نمایی", "—", "—", "—"],
        ["آلوارز [۶]", "Lorenz سه‌بعدی", "بله", "معادلات لورنز", "تمرکز بر تصاویر رنگی", "7.9974", "99.65", "30.35"],
        ["خدادادی [۷]", "Chen + Logistic", "بله", "ترکیب کانال‌ها", "دو سیستم آشوبی متوالی", "7.9971", "99.58", "33.21"],
        ["وانگ و لی [۸]", "Logistic", "بله", "XOR", "سادگی پیاده‌سازی", "7.9968", "99.57", "33.38"],
        ["مقاله [۹]", "نگاشت‌های غیرخطی", "بله", "XOR دینامیکی", "پردازش مستقل کانال‌ها", "—", "—", "—"],
        ["لی [۱۰]", "LEL نمایی", "—", "—", "توالی پرپیچیدگی بالا", "—", "—", "—"],
        ["صادقیان [۱۱]", "آشوب + GA + PSO", "—", "—", "بهینه‌سازی پارامتر", "—", "—", "—"],
        ["فوق‌آشوب [۱۲]", "Hyper-chaos", "بله", "زنجیره‌ای", "سرعت بالا (۰.۰۲۳ s)", "7.9990", "99.61", "33.46"],
        ["مقاله [۱۳]", "Lorenz+Henon+2D-Log", "بله", "جداگانه per کانال", "ترکیب سه سیستم", "—", "—", "—"],
        ["روش پیشنهادی", "Chen + ۹ ECM", "Chen (RK45)", "XOR پویا", "۹ ECM per پیکسل", "8.0", "95.13", "30.14"],
    ]

    table = add_table_after(cap, len(rows_data) + 1, len(headers))
    for j, htext in enumerate(headers):
        set_cell(table.rows[0].cells[j], htext, bold=True)
    for i, row in enumerate(rows_data, start=1):
        for j, val in enumerate(row):
            set_cell(table.rows[i].cells[j], val, bold=(row[0] == "روش پیشنهادی"))

    tbl_el = cap._p.getnext()
    note_p = OxmlElement("w:p")
    tbl_el.addnext(note_p)
    note = Paragraph(note_p, cap._parent)
    note.add_run(
        "توضیح: مقادیر آنتروپی، NPCR و UACI از گزارش مقالات مرجع یا از ارزیابی مشابه (جدول ۴-۷ برای روش پیشنهادی) استخراج شده‌اند. «—» به معنای عدم گزارش صریح آن معیار در منبع است."
    )

    for p in doc.paragraphs:
        if p.text.strip().startswith("جدول2-2"):
            insert_paragraph_after(
                p,
                "جدول2-3. مقایسه تطبیقی روش‌های رمزنگاری تصویر رنگی مبتنی بر آشوب\t57",
            )
            break

    # TOC entry for 2-6-1
    for p in doc.paragraphs:
        if p.text.strip().startswith("2-6.") and "پیشینه" in p.text:
            insert_paragraph_after(p, "2-6-1. مقایسه تطبیقی پژوهش‌های مرجع\t57")
            break

    doc.save(DOCX)
    print(f"Added literature comparison to {DOCX.name}")


if __name__ == "__main__":
    main()
