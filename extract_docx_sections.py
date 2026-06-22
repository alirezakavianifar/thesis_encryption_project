"""Extract section headings and literature review content from docx."""
from docx import Document
from pathlib import Path

doc = Document(Path(r"e:\projects\thesis_project_v2\main_updated.docx"))
out = Path(r"e:\projects\thesis_project_v2\doc_sections.txt")

lines = []
for i, p in enumerate(doc.paragraphs):
    t = p.text.strip()
    if not t:
        continue
    style = p.style.name if p.style else ""
    # headings or section markers
    if (
        "Heading" in style
        or t.startswith("فصل")
        or "مرور" in t
        or "ادبیات" in t
        or t.startswith("2-")
        or "مطالعات" in t
        or len(t) < 80 and any(c.isdigit() for c in t[:5])
    ):
        lines.append(f"[{i}|{style}] {t}")

with open(out, "w", encoding="utf-8") as f:
    f.write("\n".join(lines))
    f.write("\n\n=== FULL PARAS 80-200 ===\n")
    for i, p in enumerate(doc.paragraphs):
        if 80 <= i <= 250:
            t = p.text.strip()
            if t:
                f.write(f"[{i}] {t}\n\n")

print(f"Wrote {out}")
