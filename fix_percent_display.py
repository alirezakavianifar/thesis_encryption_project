"""Fix reversed/misformatted percentages and numeric phrases in thesis docx."""
import re
import shutil
import tempfile
import zipfile
from pathlib import Path

from docx import Document

DOCX_FILES = [
    Path(r"e:\projects\thesis_project_v2\main_updated.docx"),
    Path(r"e:\projects\thesis_project_v2\main_updated_corrected.docx"),
]

LRM = "\u200e"


def replace_in_text(text: str) -> tuple[str, int]:
    n = 0
    replacements = [
        ("NPCR > 99. 6% ", "NPCR > 99.60% "),
        ("NPCR &gt; 99. 6% ", "NPCR &gt; 99.60% "),
        ("UACI ≈ 33. 46% ", "UACI ≈ 33.46% "),
        ("(۳۳.۴۶%)", "(۳۳.۴۶٪)"),
        ("33. 46%", "33.46%"),
        ("99. 6%", "99.60%"),
        (
            "عبارتند از:  99.60% > NPCR و 33.46% ≈ UACI",
            "عبارتند از: NPCR > 99.60% و UACI ≈ 33.46%",
        ),
        (
            "عبارتند از:  99.60% &gt; NPCR و 33.46% ≈ UACI",
            "عبارتند از: NPCR &gt; 99.60% و UACI ≈ 33.46%",
        ),
    ]
    for old, new in replacements:
        if old in text:
            text = text.replace(old, new)
            n += 1

    bidi_phrases = [
        (
            "همبستگی کاهش از ۰.۹۹۸۹ به ۰.۰۱۲۶",
            f"همبستگی کاهش از {LRM}۰.۹۹۸۹{LRM} به {LRM}۰.۰۱۲۶",
        ),
        (
            "کاهش همبستگی از ۰.۹۹۸۹ به ۰.۰۱۲۶ (کاهش ۹۸.۷٪)",
            f"کاهش همبستگی از {LRM}۰.۹۹۸۹{LRM} به {LRM}۰.۰۱۲۶{LRM} (کاهش {LRM}۹۸.۷٪)",
        ),
        (
            "از میانگین ۰.۹۹۸۹ در تصاویر اصلی به ۰.۰۱۲۶ در تصاویر رمزنگاری‌شده",
            f"از میانگین {LRM}۰.۹۹۸۹{LRM} در تصاویر اصلی به {LRM}۰.۰۱۲۶{LRM} در تصاویر رمزنگاری‌شده",
        ),
        (
            "کاهش میانگین از ۰.۹۹۸۹ به ۰.۰۱۲۶",
            f"کاهش میانگین از {LRM}۰.۹۹۸۹{LRM} به {LRM}۰.۰۱۲۶",
        ),
        (
            "برابر با ۰.۰۱۲۶ است، در مقابل میانگین ۰.۹۹۸۹",
            f"برابر با {LRM}۰.۰۱۲۶{LRM} است، در مقابل میانگین {LRM}۰.۹۹۸۹",
        ),
        (
            "بیش از ۸۲٪ پیکسل",
            f"بیش از {LRM}۸۲٪{LRM} پیکسل",
        ),
        (
            "حداکثر ۹۹.۴۸٪ برای Peppers",
            f"حداکثر {LRM}۹۹.۴۸٪{LRM} برای Peppers",
        ),
        (
            "بیش از ۸۲٪ (و در سایر تصاویر بیش از ۹۸٪)",
            f"بیش از {LRM}۸۲٪{LRM} (و در سایر تصاویر بیش از {LRM}۹۸٪{LRM})",
        ),
        (
            "بین ۲۶.۳۴٪ (Airplane) تا ۳۲.۹۱٪ (Peppers) متغیر است (با میانگین کل ۳۰.۱۹٪)",
            f"بین {LRM}۲۶.۳۴٪{LRM} (Airplane) تا {LRM}۳۲.۹۱٪{LRM} (Peppers) متغیر است (با میانگین کل {LRM}۳۰.۱۹٪{LRM})",
        ),
        (
            "میانگین آنتروپی ۷.۹۹۹۳ بیت (۰.۰۰۹٪ انحراف از ایده‌آل)",
            f"میانگین آنتروپی {LRM}۷.۹۹۹۳{LRM} بیت ({LRM}۰.۰۰۹٪{LRM} انحراف از ایده‌آل)",
        ),
        ("NPCR تا ۹۹.۴۸٪؛", f"NPCR تا {LRM}۹۹.۴۸٪{LRM}؛"),
        (
            "برای تصویر Peppers به ۹۹.۴۸٪ نزدیک",
            f"برای تصویر Peppers به {LRM}۹۹.۴۸٪{LRM} نزدیک",
        ),
        (
            "NPCR تا ۹۹.۴۸٪ برای تصویر Peppers و میانگین ۹۴.۹۲٪",
            f"NPCR تا {LRM}۹۹.۴۸٪{LRM} برای تصویر Peppers و میانگین {LRM}۹۴.۹۲٪",
        ),
        (
            "مقادیر تصویر Airplane (۲۶.۳۴٪) فاصله معناداری از مقدار ایده‌آل (۳۳.۴۶٪)",
            f"مقادیر تصویر Airplane ({LRM}۲۶.۳۴٪{LRM}) فاصله معناداری از مقدار ایده‌آل ({LRM}۳۳.۴۶٪{LRM})",
        ),
        (
            "مقدار UACI برای تصویر Airplane (۲۶.۳۴٪) اندکی از مقدار ایده‌آل (۳۳.۴۶٪)",
            f"مقدار UACI برای تصویر Airplane ({LRM}۲۶.۳۴٪{LRM}) اندکی از مقدار ایده‌آل ({LRM}۳۳.۴۶٪{LRM})",
        ),
        (
            "یعنی کاهشی معادل ۹۸.۷٪",
            f"یعنی کاهشی معادل {LRM}۹۸.۷٪",
        ),
        (
            "معادل کاهش ۹۸.۷٪ است",
            f"معادل کاهش {LRM}۹۸.۷٪{LRM} است",
        ),
    ]
    for old, new in bidi_phrases:
        if old in text and new not in text:
            text = text.replace(old, new)
            n += 1
    return text, n


def set_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.text = new_text


def fix_document(path: Path) -> int:
    doc = Document(path)
    total = 0

    for para in doc.paragraphs:
        new, n = replace_in_text(para.text)
        if n:
            set_paragraph_text(para, new)
            total += n

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    new, n = replace_in_text(para.text)
                    if n:
                        set_paragraph_text(para, new)
                        total += n

    doc.save(path)
    return total


if __name__ == "__main__":
    for path in DOCX_FILES:
        if path.exists():
            changes = fix_document(path)
            print(f"Fixed {path.name}: {changes} paragraph/cell updates")
